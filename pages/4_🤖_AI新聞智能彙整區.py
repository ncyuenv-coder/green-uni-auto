import streamlit as st
import pandas as pd
import gspread
import datetime
import io
import requests
import re
import time
import random
import urllib3
import base64
from bs4 import BeautifulSoup
from PIL import Image
import google.generativeai as genai
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import qn

# 🌟 關閉 SSL 憑證驗證警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ==========================================
# 🌟 全域設定 & 資安防護罩
# ==========================================
if st.session_state.get("authentication_status") is not True:
    st.warning("⚠️ 請先至首頁登入系統！"); st.stop()

# 🛡️ 僅限 admin_ui 檢視與管理
if st.session_state.get("username") != "admin_ui":
    st.error("🚫 權限不足！此頁面僅限系統管理員 (admin_ui) 存取。")
    st.stop()

st.set_page_config(page_title="嘉大 AI 新聞智能彙整", page_icon="📰", layout="wide")

# 📂 新聞照片存放的 Google Drive 資料夾 ID
NEWS_IMG_FOLDER_ID = "1VNOna4gRdtTIiFPc4XqMJU2jP01LcyXM" 

# 🌟 初始化 Gemini AI 大腦
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=GEMINI_API_KEY)
    # ⚡ 替換為專為高頻率、文字摘要打造的輕量化模型，避免 429 錯誤且速度更快
    ai_model = genai.GenerativeModel('gemini-1.5-flash-8b') 
except Exception as e:
    st.error(f"⚠️ 無法載入 Gemini API Key，或設定有誤：{e}")
    st.stop()

st.markdown("""
<style>
    button[data-baseweb="tab"] { background-color: #E6F0F9 !important; border-radius: 8px 8px 0px 0px !important; margin-right: 4px !important; padding: 10px 20px !important; border: 1px solid #AED6F1 !important; border-bottom: none !important; transition: all 0.3s ease; }
    button[data-baseweb="tab"] p { font-size: 1.25em !important; font-weight: bold !important; color: #2C3E50 !important; }
    button[data-baseweb="tab"][aria-selected="true"] { background-color: #154360 !important; border: 1px solid #0B2331 !important; border-bottom: 3px solid #0B2331 !important; }
    button[data-baseweb="tab"][aria-selected="true"] p { color: #FFFFFF !important; }
    .morandi-title { background-color: #738A96; color: white; padding: 15px 20px; border-radius: 8px; font-weight: bold; font-size: 1.5em; margin-bottom: 20px; text-align: center; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    .morandi-dark-title { background-color: #5C6B73; color: white; padding: 12px 20px; border-radius: 6px; font-weight: bold; font-size: 1.3em; margin-bottom: 10px; margin-top: 15px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
    div.stButton > button[kind="primary"] { border-radius: 8px !important; font-weight: bold !important; font-size: 1.2em !important; padding: 10px 20px !important; background-color: #D4A373 !important; border: none !important; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    div.stButton > button[kind="secondary"] { border-radius: 8px !important; font-weight: bold !important; font-size: 1.2em !important; padding: 10px 20px !important; background-color: #8FAAB8 !important; color: white !important; border: none !important; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    div.stButton > button[kind="secondary"]:hover { background-color: #738A96 !important; }
    [data-testid="stDownloadButton"] button { background-color: #D4A373 !important; color: white !important; border: none !important; border-radius: 8px !important; font-weight: bold !important; font-size: 1.3em !important; padding: 12px 24px !important; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    .raw-box { background-color: #FDF6E3; padding: 20px; border-left: 5px solid #E6C27A; border-radius: 6px; height: 500px; overflow-y: auto; line-height: 1.8; font-size: 1.1em; color: #333;}
    .action-box { background-color: #F4F6F6; padding: 20px; border-radius: 8px; border: 1px solid #D5DBDB; margin-top: 15px; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 🔌 資料庫連線與精準讀寫引擎
# ==========================================
@st.cache_resource
def get_gcp_credentials():
    sk = st.secrets["gcp_oauth"].to_dict()
    return Credentials(token=None, refresh_token=sk["refresh_token"], token_uri="https://oauth2.googleapis.com/token", client_id=sk["client_id"], client_secret=sk["client_secret"])

def load_sheet(sheet_name, max_retries=3):
    for attempt in range(max_retries):
        try:
            return pd.DataFrame(gspread.authorize(get_gcp_credentials()).open_by_key('1JNbpZoZHWZRrIzn0whcQFnCDkOZghZmMyFidLE7dxT8').worksheet(sheet_name).get_all_records()).rename(columns=lambda x: str(x).strip())
        except Exception:
            if attempt < max_retries - 1: time.sleep(2)
            else: return pd.DataFrame()

def save_raw_to_db(ws, record, max_retries=5):
    for attempt in range(max_retries):
        try:
            row = [
                datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                str(record['題號']).strip(), str(record['中文標題']).strip(), record['新聞日期'],
                record['新聞標題'], record['原始內容'], "", 
                ",".join(record['照片清單']), record['新聞連結']
            ]
            ws.append_row(row)
            return True
        except Exception:
            if attempt < max_retries - 1: time.sleep(random.uniform(2.0, 5.0))
            else: return False

def delete_rows_from_db(ws, row_indices, max_retries=3):
    for attempt in range(max_retries):
        try:
            for r_idx in sorted(row_indices, reverse=True):
                # 強制轉換 r_idx 為原生 int，解決 numpy.int64 導致的序列化失敗問題
                ws.delete_row(int(r_idx))
            return True
        except Exception as e:
            if attempt < max_retries - 1: time.sleep(2)
            else: 
                st.error(f"Sheet 刪除失敗: {e}")
                return False

def update_q_ids_in_db(ws, updates, max_retries=3):
    for attempt in range(max_retries):
        try:
            headers = [str(h).strip() for h in ws.get_all_values()[0]]
            q_id_col = headers.index('對應題號') + 1
            q_title_col = headers.index('中文標題') + 1
            cells = []
            for u in updates:
                # 強制轉換 row 和 col 為原生 int
                cells.append(gspread.Cell(row=int(u['row']), col=int(q_id_col), value=str(u['new_id'])))
                cells.append(gspread.Cell(row=int(u['row']), col=int(q_title_col), value=str(u['new_title'])))
            ws.update_cells(cells)
            return True
        except Exception:
            if attempt < max_retries - 1: time.sleep(2)
            else: return False

def update_ai_summary_by_row(ws, row_idx, ai_col_idx, ai_summary, max_retries=5):
    for attempt in range(max_retries):
        try:
            # 強制轉換 row 和 col 為原生 int
            ws.update_cell(int(row_idx), int(ai_col_idx), str(ai_summary))
            return True
        except Exception:
            if attempt < max_retries - 1: time.sleep(random.uniform(2.0, 5.0))
            else: return False

# ==========================================
# 🌟 Drive 讀取、寫入與刪除引擎
# ==========================================
def delete_drive_files(file_ids, max_retries=3):
    if not file_ids: return
    try:
        creds = get_gcp_credentials()
        drive_service = build('drive', 'v3', credentials=creds)
        for file_id in file_ids:
            fid = file_id.strip()
            if not fid: continue
            for attempt in range(max_retries):
                try:
                    drive_service.files().delete(fileId=fid).execute()
                    break
                except Exception:
                    if attempt < max_retries - 1: time.sleep(1)
    except Exception:
        pass 

def drive_id_to_bytes(file_id):
    try:
        creds = get_gcp_credentials()
        drive_service = build('drive', 'v3', credentials=creds)
        request = drive_service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done: status, done = downloader.next_chunk()
        fh.seek(0)
        return fh
    except: return None

def get_drive_image_b64(file_id):
    fh = drive_id_to_bytes(file_id)
    if fh:
        return base64.b64encode(fh.read()).decode('utf-8')
    return ""

# ==========================================
# 🕸️ 爬蟲引擎與圖片雙重瘦身上傳模組 
# ==========================================
def download_compress_upload_image(img_url, drive_service, file_name_prefix, max_retries=3):
    for attempt in range(max_retries):
        try:
            headers = {'User-Agent': 'Mozilla/5.0'}
            req = requests.get(img_url, headers=headers, timeout=15, verify=False)
            req.raise_for_status()
            img_bytes = req.content

            img = Image.open(io.BytesIO(img_bytes))
            if img.mode in ("RGBA", "P"): img = img.convert("RGB")
                
            MAX_SIZE = (1280, 1280)
            img.thumbnail(MAX_SIZE, Image.Resampling.LANCZOS)

            out_io = io.BytesIO()
            img.save(out_io, format='JPEG', optimize=True, quality=85)
            out_io.seek(0)

            if len(out_io.getvalue()) > len(img_bytes):
                out_io = io.BytesIO(img_bytes)
                mime_type = req.headers.get('Content-Type', 'image/jpeg')
            else:
                mime_type = 'image/jpeg'

            file_meta = {'name': f"{file_name_prefix}.jpg", 'parents': [NEWS_IMG_FOLDER_ID]}
            media = MediaIoBaseUpload(out_io, mimetype=mime_type, resumable=True)
            file = drive_service.files().create(body=file_meta, media_body=media, fields='id').execute()
            return file.get('id')
        except Exception:
            if attempt < max_retries - 1: time.sleep(random.uniform(2.0, 4.0))
            else: return None

def get_news_list(start_date, end_date):
    base_url = "https://www.ncyu.edu.tw"
    headers = {'User-Agent': 'Mozilla/5.0'}
    news_list = []
    seen_links = set()
    page = 1
    old_news_count = 0 
    
    while page <= 40: 
        list_url = f"{base_url}/ncyu/Subject?nodeId=835&page={page}"
        success = False
        
        for attempt in range(3):
            try:
                req = requests.get(list_url, headers=headers, timeout=15, verify=False)
                req.encoding = 'utf-8'
                success = True
                break
            except Exception:
                time.sleep(2)
                
        if not success: break
        
        try:
            soup = BeautifulSoup(req.text, 'html.parser')
            links = soup.find_all('a', href=True)
            has_news_in_page = False
            
            for a in links:
                href = a['href']
                if 'Subject/Detail/' in href and 'nodeId=835' in href:
                    has_news_in_page = True
                    full_link = href if href.startswith('http') else (base_url + href if href.startswith('/') else base_url + "/ncyu/" + href)
                    if full_link in seen_links: continue
                    seen_links.add(full_link)
                    
                    node = a
                    date_str = ""
                    for _ in range(4):
                        if not node.parent: break
                        text = node.parent.get_text(separator=' ')
                        date_match = re.search(r'(20\d{2})\s*[-/.]\s*(\d{1,2})\s*[-/.]\s*(\d{1,2})', text)
                        if date_match:
                            y, m, d = date_match.groups()
                            date_str = f"{y}-{int(m):02d}-{int(d):02d}"
                            break
                        node = node.parent
                    
                    if date_str:
                        news_date = datetime.datetime.strptime(date_str, "%Y-%m-%d").date()
                        if news_date < start_date: old_news_count += 1
                        else: old_news_count = 0 
                        
                        if start_date <= news_date <= end_date:
                            title = a.get('title', '').strip() or a.text.strip()
                            title = re.sub(r'\s+', ' ', title).strip() 
                            news_list.append({"新聞日期": date_str, "新聞標題": title, "新聞連結": full_link})
                            
            if not has_news_in_page: break 
            if old_news_count > 25: break 
        except Exception: break
        
        page += 1
        time.sleep(0.5)
    return news_list

def get_news_content(url, max_retries=3):
    base_url = "https://www.ncyu.edu.tw"
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    for attempt in range(max_retries):
        try:
            d_req = requests.get(url, headers=headers, timeout=15, verify=False)
            d_req.encoding = 'utf-8'
            d_soup = BeautifulSoup(d_req.text, 'html.parser')
            
            for tag in d_soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
                tag.decompose()
                
            noise_patterns = re.compile(r'menu|nav|footer|breadcrumb|sidebar|top-bar|bottom-bar|search|language|header', re.I)
            for div in d_soup.find_all(['div', 'ul', 'ol', 'section']):
                class_str = ' '.join(div.get('class', [])) if isinstance(div.get('class'), list) else str(div.get('class', ''))
                id_str = str(div.get('id', ''))
                if noise_patterns.search(class_str) or noise_patterns.search(id_str):
                    div.decompose()

            specific_classes = ['m-edit', 'user_edit', 'article-content', 'news-content', 'CCms_Content', 'cg-desc', 'editor', 'app-article']
            content_div = None
            for cls in specific_classes:
                content_div = d_soup.find('div', class_=re.compile(cls, re.I))
                if content_div: break
                
            if not content_div:
                best_node = None
                max_score = 0
                for node in d_soup.find_all(['div', 'article', 'main']):
                    text_len = len(node.get_text(strip=True))
                    if text_len < 50: continue
                    link_len = sum(len(a.get_text(strip=True)) for a in node.find_all('a'))
                    if text_len > 0 and (link_len / text_len) > 0.4: continue 
                    
                    p_count = len(node.find_all('p', recursive=False))
                    score = text_len + (p_count * 50)
                    if score > max_score:
                        max_score = score
                        best_node = node
                content_div = best_node if best_node else d_soup.find('body')

            for br in content_div.find_all(['br', 'hr']):
                br.replace_with('\n')
            for block in content_div.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'tr']):
                block.append('\n')

            full_text = content_div.get_text(separator=' ', strip=True)
            
            garbage_words = [
                '國立嘉義大學', ':::', '回首頁', '網站導覽', '分眾導覽', '新生教務專欄', 
                '學生', '教師', '職員工', '校友', '民眾', '聯絡我們', 'ENGLISH', 
                'Select Language', '中文 (繁體)', '中文 (简体)', '日本語', 'Bahasa Indonesia', 
                'हिन्दी', 'Español', 'বাংলা', 'Français', 'العربية', 'English', 'ไทย', 
                'اردو', 'Bahasa Melayu', 'Filipino', 'Tiếng Việt', 'မြန်မာ', '한국어', 
                '請輸入關鍵字', '搜尋', '友善列印(開新視窗)', '分享至臉書(開新視窗)', '分享至Line(開新視窗)', '回上一頁'
            ]
            
            clean_lines = []
            for line in full_text.split('\n'):
                line_str = line.strip()
                line_str = re.sub(r' +', ' ', line_str) 
                line_str = re.sub(r'[（\(][^）\)]*(照片|攝影|拍攝|提供)[^）\)]*[）\)]', '', line_str)
                line_str = re.sub(r'(?:^|\s)圖\s*\d+\s*[：:].*?(?:[）\)]|。|$)', '', line_str)
                line_str = line_str.strip()
                
                if not line_str: continue
                if line_str in garbage_words: continue 
                if len(line_str) <= 3 and '::' in line_str: continue
                clean_lines.append(line_str)
                
            final_text = '\n'.join(clean_lines)
            
            img_urls = []
            for img in content_div.find_all('img'):
                src = img.get('src')
                if src and not src.startswith('data:'):
                    img_link = src if src.startswith('http') else base_url + src
                    if 'icon' not in img_link.lower() and 'logo' not in img_link.lower() and 'banner' not in img_link.lower():
                        img_urls.append(img_link)
                        
            return {"原始內容": final_text[:3000], "照片清單": img_urls}
        except Exception:
            if attempt < max_retries - 1: time.sleep(2)
            
    return {"原始內容": "無法擷取內文", "照片清單": []}

# ==========================================
# 🧠 Gemini AI 摘要引擎 
# ==========================================
def process_news_with_ai(news_item, q_title):
    prompt = f"""
    你現在是大學永續發展(SDGs)評比專家。
    這篇新聞已經被確認可能符合評比題目『{q_title}』。請根據以下【新聞內容】進行濃縮。
    
    【新聞標題】：{news_item['新聞標題']}
    【新聞內容】：{news_item['原始內容']}
    
    任務要求：
    1. 將新聞濃縮精華為 150~300 字，重點放在與「{q_title}」相關的行動或成效。
    2. 濃縮內容若有列點，請使用「1. 」「2. 」的格式。
    3. 請在摘要的最下方，獨立一行，明確標註這篇新聞對應的 SDGs 指標 (例如：【對應SDGs】：SDG 4 優質教育)。
    """
    try:
        response = ai_model.generate_content(prompt)
        try:
            return response.text.strip()
        except ValueError:
            return f"⚠️ 內容可能因觸發 Google 安全機制（如特定敏感詞彙）而被拒絕摘要。"
    except Exception as e: 
        return f"❌ 系統錯誤: {str(e)}"

# ==========================================
# 🌟 Word 報表產製模組
# ==========================================
def format_report_text_to_html(text):
    text = str(text)
    text = re.sub(r'(^|\n)(\d+[\.\)]|[-•*])\s*\n\s*', r'\1\2 ', text)
    html = ""
    for line in text.split('\n'):
        line = line.strip()
        if not line: html += "<div style='height: 10px;'></div>"; continue
        match = re.match(r'^([0-9a-zA-Z]+[\.、\)]|[\(（][0-9a-zA-Z一二三四五六七八九十]+[\)）]|[一二三四五六七八九十]+、|[-•*])\s*(.*)', line)
        if match: 
            marker = match.group(1); content = match.group(2)
            html += f"<div style='padding-left: 2.2em; text-indent: -2.2em; margin-bottom: 5px; line-height: 1.6;'>{marker} {content}</div>"
        else: html += f"<div style='margin-bottom: 5px; line-height: 1.6;'>{line}</div>"
    return html.replace('\n', ' ')

def set_run_font(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def generate_ai_word_report(q_id, q_title, news_records):
    doc = Document()
    doc.add_heading(f'{q_id} {q_title}', level=1)
    for idx, record in enumerate(news_records):
        if idx > 0: doc.add_paragraph("="*40)
        doc.add_heading(f"新聞標題：{record['新聞標題']}", level=2)
        p_date = doc.add_paragraph(); p_date.add_run(f"發布日期：{record['新聞日期']}").italic = True
        
        file_ids = str(record['照片清單']).split(',') if record['照片清單'] else []
        file_ids = [fid.strip() for fid in file_ids if fid.strip()]
        
        if file_ids:
            table = doc.add_table(rows=0, cols=2); table.style = 'Table Grid'
            for i in range(0, len(file_ids), 2):
                row = table.add_row()
                c1 = row.cells[0]; c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_b1 = drive_id_to_bytes(file_ids[i])
                if img_b1: 
                    try: p1.add_run().add_picture(img_b1, width=Cm(7.5))
                    except: p1.add_run("(圖檔格式不支援)")
                else: p1.add_run("(圖載入失敗)")
                
                c2 = row.cells[1]; c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i + 1 < len(file_ids):
                    img_b2 = drive_id_to_bytes(file_ids[i+1])
                    if img_b2: 
                        try: p2.add_run().add_picture(img_b2, width=Cm(7.5))
                        except: p2.add_run("(圖檔格式不支援)")
                    else: p2.add_run("(圖載入失敗)")
                    
        doc.add_paragraph()
        for line in str(record['AI摘要']).replace('\r', '').split('\n'):
            line = line.strip()
            if not line: continue
            p = doc.add_paragraph()
            match = re.match(r'^([0-9a-zA-Z]+[\.、\)]|[\(（][0-9a-zA-Z一二三四五六七八九十]+[\)）]|[一二三四五六七八九十]+、|[-•*])\s*(.*)', line)
            if match:
                marker = match.group(1); content = match.group(2)
                p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.first_line_indent = Cm(-0.8)
                p.add_run(f"{marker} {content}")
            else: p.add_run(line)
        p_link = doc.add_paragraph(); p_link.add_run("原始新聞連結: ").bold = True; p_link.add_run(record['新聞連結'])
        
    for p in doc.paragraphs:
        if not p.style.name.startswith('Heading'): p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs: set_run_font(r)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs: set_run_font(r)
    out = io.BytesIO(); doc.save(out); out.seek(0)
    return out

# ==========================================
# 🚀 網頁主畫面與 UI 邏輯
# ==========================================
st.markdown("<div class='morandi-title'>🤖 綠色大學 AI 新聞智能彙整中心</div>", unsafe_allow_html=True)

tab_scrape, tab_ai, tab_view = st.tabs(["🕷️ 階段一：爬蟲抓取", "🔍 階段二：審核與 AI 改寫", "📥 階段三：檢視與下載"])

# ==========================================
# 🕷️ 階段一：爬蟲抓取與入庫
# ==========================================
with tab_scrape:
    st.markdown("<div class='morandi-dark-title'>⚙️ 第一階段：爬取原始新聞與入庫</div>", unsafe_allow_html=True)
    c_start, c_end, c_btn = st.columns([2, 2, 2])
    with c_start: start_date = st.date_input("新聞起日", datetime.date(2024, 1, 1))
    with c_end: end_date = st.date_input("新聞迄日", datetime.date.today())
    
    st.info("💡 **【強固防護版】** 系統已啟動「自動斷線重連」、「防 API 擁塞限制」及「智慧跳過已存在新聞」等機制，可安心抓取大量新聞。")
    
    if st.button("🕷️ 開始啟動關鍵字爬取與照片儲存", type="secondary", use_container_width=True):
        with st.spinner("🤖 正在掃描並準備資料庫連線，請稍候..."):
            df_targets = load_sheet("原始新聞抓取")
            df_existing = load_sheet("AI新聞資料庫")
            
            existing_set = set()
            if not df_existing.empty and '對應題號' in df_existing.columns and '新聞連結' in df_existing.columns:
                for _, r in df_existing.iterrows():
                    row_id = str(r.get('對應題號','')).strip()
                    ids = [x.strip() for x in re.split(r'[,、/，\s]+', row_id) if x.strip()]
                    link = str(r.get('新聞連結','')).strip()
                    for i in ids:
                        existing_set.add(f"{i}_{link}")
            
            if df_targets.empty:
                st.error("❌ 找不到《原始新聞抓取》工作表，或表內無資料！")
            else:
                st.toast("🕷️ 正在地毯式掃描新聞列表...", icon="👀")
                basic_news_list = get_news_list(start_date, end_date)
                
                if not basic_news_list:
                    st.warning(f"在 {start_date} 至 {end_date} 區間內未抓取到任何新聞。")
                else:
                    st.toast(f"✅ 成功獲取 {len(basic_news_list)} 篇新聞清單，準備比對...", icon="🎯")
                    
                    matched_tasks = []
                    skipped_count = 0
                    kw_col = '搜尋關鍵字或判斷準則' if '搜尋關鍵字或判斷準則' in df_targets.columns else '關鍵字或判斷準則'
                    
                    for _, target in df_targets.iterrows():
                        q_id = str(target.get('題號', '')).strip()
                        q_title = str(target.get('中文標題', '')).strip()
                        raw_kw = str(target.get(kw_col, ''))
                        if not raw_kw.strip() or raw_kw.lower() == 'nan': raw_kw = q_title
                        keywords = [k.strip() for k in re.split(r'[、,，]', raw_kw) if k.strip()]
                        
                        if not q_id: continue
                        
                        for news in basic_news_list:
                            clean_title = re.sub(r'\s+', '', news['新聞標題']).lower()
                            for kw in keywords:
                                clean_kw = re.sub(r'\s+', '', kw).lower()
                                if clean_kw and clean_kw in clean_title:
                                    check_key = f"{q_id}_{news['新聞連結']}"
                                    if check_key not in existing_set:
                                        matched_tasks.append({'q_id': q_id, 'q_title': q_title, 'news': news})
                                        existing_set.add(check_key)
                                    else:
                                        skipped_count += 1
                                    break
                    
                    if not matched_tasks:
                        st.success(f"🎉 所有符合關鍵字的新聞都已經在資料庫中了 (本次自動跳過了 {skipped_count} 筆已存在的紀錄)！")
                    else:
                        st.info(f"💡 共比對出 {len(matched_tasks) + skipped_count} 筆新聞，自動跳過 {skipped_count} 筆已存在的資料，實際將擷取並上傳 {len(matched_tasks)} 筆。")
                        
                        success_count = 0
                        total_tasks = len(matched_tasks)
                        progress_text = st.empty()
                        my_bar = st.progress(0)
                        
                        drive_service = build('drive', 'v3', credentials=get_gcp_credentials())
                        gc = gspread.authorize(get_gcp_credentials())
                        ws_ai_db = gc.open_by_key('1JNbpZoZHWZRrIzn0whcQFnCDkOZghZmMyFidLE7dxT8').worksheet("AI新聞資料庫")
                        
                        for i, task in enumerate(matched_tasks):
                            q_id = task['q_id']
                            news_title = task['news']['新聞標題']
                            
                            progress_text.markdown(f"**📥 ({i+1}/{total_tasks}) 擷取與上傳中：** 題目 {q_id} v.s. 「{news_title}」")
                            my_bar.progress((i + 1) / total_tasks)
                            
                            detail = get_news_content(task['news']['新聞連結'])
                            full_news = {**task['news'], **detail}
                            
                            uploaded_file_ids = []
                            for img_idx, img_url in enumerate(full_news['照片清單']):
                                clean_name_prefix = f"News_{q_id}_{re.sub(r'[/\\:*?\"<>|]', '', news_title)[:15]}_{img_idx+1}"
                                fid = download_compress_upload_image(img_url, drive_service, clean_name_prefix)
                                if fid: uploaded_file_ids.append(fid)
                            
                            record = {
                                '題號': q_id, '中文標題': task['q_title'], '新聞日期': full_news['新聞日期'],
                                '新聞標題': full_news['新聞標題'], '原始內容': full_news['原始內容'],
                                '照片清單': uploaded_file_ids, '新聞連結': full_news['新聞連結']
                            }
                            
                            if save_raw_to_db(ws_ai_db, record):
                                success_count += 1
                            time.sleep(1.0) 
                                    
                        progress_text.empty()
                        my_bar.empty()
                        st.success(f"🎉 爬蟲作業穩健完成！成功新增擷取 {success_count} 筆新聞，請前往「階段二」進行人工審核與改寫。")

# ==========================================
# 🔍 階段二：單一新聞獨立處理模式 (UI 全面升級)
# ==========================================
with tab_ai:
    st.markdown("<div class='morandi-dark-title'>🔍 第二階段：人工審核與 AI 智慧改寫</div>", unsafe_allow_html=True)
    
    df_ai_db = load_sheet("AI新聞資料庫")
    df_targets = load_sheet("原始新聞抓取")
    
    required_cols = ['對應題號', '中文標題', '新聞日期', '新聞標題', '原始內容', 'AI摘要', '照片清單', '新聞連結']
    
    if df_ai_db.empty:
        st.warning("目前資料庫為空，請先至「階段一」進行爬蟲抓取。")
    elif not all(col in df_ai_db.columns for col in required_cols):
        missing = [col for col in required_cols if col not in df_ai_db.columns]
        st.error(f"⚠️ 您的《AI新聞資料庫》缺少必要的欄位標題：`{', '.join(missing)}`")
    else:
        # 建立有效題號清單 (給修改題號的選單使用)
        valid_q_list = []
        if not df_targets.empty:
            for _, r in df_targets.iterrows():
                qid = str(r.get('題號', '')).strip()
                qtitle = str(r.get('中文標題', '')).strip()
                if qid:
                    valid_q_list.append(f"{qid} - {qtitle}")
        
        # 篩選出尚未進行 AI 摘要的新聞
        df_ai_db['AI摘要'] = df_ai_db['AI摘要'].astype(str).str.strip()
        df_pending = df_ai_db[df_ai_db['AI摘要'] == ''].copy()
        
        if df_pending.empty:
            st.success("🎉 太棒了！資料庫中所有新聞都已經完成 AI 摘要改寫囉！")
        else:
            # 建立選單選項文字
            df_pending['_row_idx'] = df_pending.index + 2
            df_pending['對應題號'] = df_pending['對應題號'].astype(str).str.strip() + " - " + df_pending['中文標題'].astype(str).str.strip()
            df_pending['下拉選項'] = "【" + df_pending['對應題號'].astype(str) + "】 " + df_pending['新聞標題'].astype(str)
            
            st.markdown(f"#### 📝 尚有 <span style='color:red;'>{len(df_pending)}</span> 筆待處理新聞", unsafe_allow_html=True)
            st.info("💡 請從下方選單挑選一則新聞，預覽內容後在底部點擊對應的處理按鈕即可。")
            
            # 1. 獨立新聞選擇下拉選單
            selected_news_str = st.selectbox("📌 請選擇要處理的新聞：", df_pending['下拉選項'].tolist())
            
            if selected_news_str:
                # 擷取當前選擇新聞的詳細資料
                target_row = df_pending[df_pending['下拉選項'] == selected_news_str].iloc[0]
                real_row_idx = target_row['_row_idx']
                current_full_id = target_row['對應題號']
                news_title = target_row['新聞標題']
                news_content = target_row['原始內容']
                news_link = target_row['新聞連結']
                news_date = target_row['新聞日期']
                
                # 2. 顯示新聞內容與預覽
                st.markdown("---")
                st.markdown(f"### 📰 {news_title}")
                st.caption(f"📅 發布日期：{news_date} | 🎯 目前對應題目：**{current_full_id}**")
                st.markdown(f"**🔗 原文連結：** [{news_link}]({news_link})")
                st.markdown(f"<div class='raw-box'>{news_content}</div>", unsafe_allow_html=True)
                
                # 3. 三大功能操作區塊
                st.markdown("<div class='action-box'>", unsafe_allow_html=True)
                st.markdown("### 🛠️ 請選擇處理模式")
                
                col_act1, col_act2 = st.columns([1, 1], gap="large")
                
                with col_act1:
                    st.markdown("#### 🚀 模式一：AI 處理")
                    st.write("確認新聞與題目相符，直接交由 Gemini 濃縮摘要。")
                    if st.button("✨ 啟動 Gemini 智慧改寫", type="primary", use_container_width=True):
                        with st.spinner("呼叫 Gemini 處理中..."):
                            gc = gspread.authorize(get_gcp_credentials())
                            ws_ai_db = gc.open_by_key('1JNbpZoZHWZRrIzn0whcQFnCDkOZghZmMyFidLE7dxT8').worksheet("AI新聞資料庫")
                            headers = [str(h).strip() for h in ws_ai_db.get_all_values()[0]]
                            ai_col_idx = headers.index('AI摘要') + 1
                            
                            target_title = current_full_id.split(" - ", 1)[1].strip() if " - " in current_full_id else ""
                            news_item = {'新聞標題': news_title, '原始內容': news_content}
                            
                            ai_summary = process_news_with_ai(news_item, target_title)
                            
                            # 攔截寫入 Google Sheet 動作，若 AI 回傳錯誤直接顯示警告
                            if "❌ 系統錯誤" in ai_summary or "⚠️" in ai_summary:
                                st.error(f"AI 處理失敗，請稍後再試：\n{ai_summary}")
                            else:
                                if update_ai_summary_by_row(ws_ai_db, real_row_idx, ai_col_idx, ai_summary):
                                    st.success("🎉 完成！成功生成新聞摘要並寫入資料庫！")
                                    # 防 429 機制，給予 API 緩衝時間
                                    time.sleep(4.5)
                                    st.rerun()
                                else:
                                    st.error("❌ 寫入資料庫失敗，請檢查網路連線。")
                                
                    st.markdown("---")
                    st.markdown("#### 🗑️ 模式二：刪除廢棄")
                    st.write("若此新聞無參考價值，將徹底從資料庫與雲端硬碟移除。")
                    if st.button("🗑️ 徹底刪除此新聞 (含照片)", type="secondary", use_container_width=True):
                        with st.spinner("正在安全刪除照片與紀錄..."):
                            gc = gspread.authorize(get_gcp_credentials())
                            ws_ai_db = gc.open_by_key('1JNbpZoZHWZRrIzn0whcQFnCDkOZghZmMyFidLE7dxT8').worksheet("AI新聞資料庫")
                            
                            # 刪除 Drive 上的照片
                            photos_str = str(target_row.get('照片清單', ''))
                            if photos_str:
                                fids = [f.strip() for f in photos_str.split(',') if f.strip()]
                                if fids: delete_drive_files(fids)
                                
                            # 刪除 Sheet 上的資料列 (加入成功與否的判斷)
                            success_del = delete_rows_from_db(ws_ai_db, [real_row_idx])
                            if success_del:
                                st.success("✅ 該篇新聞與附屬的雲端照片皆已徹底刪除！")
                                time.sleep(1.5)
                                st.rerun()
                            else:
                                st.error("❌ Google Sheet 資料刪除失敗，請再試一次。")
                            
                with col_act2:
                    st.markdown("#### ✏️ 模式三：修改題號")
                    st.write("若此新聞判斷的題目不正確，請從下方選擇正確題號並儲存。")
                    
                    try: 
                        idx_default = valid_q_list.index(current_full_id)
                    except ValueError: 
                        idx_default = 0
                        
                    new_q_selection = st.selectbox("選擇新對應題號：", valid_q_list, index=idx_default, label_visibility="collapsed")
                    
                    if st.button("💾 儲存新題號", use_container_width=True):
                        if new_q_selection != current_full_id:
                            with st.spinner("正在為您更新對應題號..."):
                                gc = gspread.authorize(get_gcp_credentials())
                                ws_ai_db = gc.open_by_key('1JNbpZoZHWZRrIzn0whcQFnCDkOZghZmMyFidLE7dxT8').worksheet("AI新聞資料庫")
                                
                                new_id = new_q_selection.split(" - ", 1)[0].strip() if " - " in new_q_selection else new_q_selection.strip()
                                new_title = new_q_selection.split(" - ", 1)[1].strip() if " - " in new_q_selection else ""
                                
                                updates = [{'row': real_row_idx, 'new_id': new_id, 'new_title': new_title}]
                                if update_q_ids_in_db(ws_ai_db, updates):
                                    st.success("✅ 題號修改成功！")
                                    time.sleep(1.5)
                                    st.rerun()
                                else:
                                    st.error("❌ 題號修改寫入失敗。")
                        else:
                            st.warning("您選擇的題號與目前相同，無須修改喔！")
                
                st.markdown("</div>", unsafe_allow_html=True)

# ==========================================
# 📥 階段三：檢視與下載彙整區
# ==========================================
with tab_view:
    st.markdown("<div class='morandi-dark-title'>📥 依題目檢視 AI 彙整成果</div>", unsafe_allow_html=True)
    
    df_ai_db = load_sheet("AI新聞資料庫")
    
    if df_ai_db.empty or not all(col in df_ai_db.columns for col in required_cols): 
        st.info("💡 目前資料庫中尚未有紀錄或欄位不齊全。")
    else:
        df_ai_db['AI摘要'] = df_ai_db['AI摘要'].astype(str).str.strip()
        df_completed = df_ai_db[df_ai_db['AI摘要'] != '']
        
        if df_completed.empty:
            st.warning("目前尚無完成 AI 摘要的資料，請先至「階段二」執行處理。")
        else:
            reported_q_ids = df_completed['對應題號'].astype(str).str.strip().unique().tolist()
            q_options = []
            for qid in reported_q_ids:
                title_match = df_completed[df_completed['對應題號'].astype(str).str.strip() == qid]
                title = title_match.iloc[0]['中文標題'] if not title_match.empty else ""
                q_options.append(f"{qid} - {title}")
                
            q_options.sort()
            sel_q = st.selectbox("選擇欲下載之題目", ["請選擇..."] + q_options, label_visibility="collapsed")
            
            if sel_q != "請選擇...":
                v_q_id = sel_q.split(" - ")[0]
                v_q_title = sel_q.split(" - ")[1]
                
                df_records = df_completed[df_completed['對應題號'].astype(str).str.strip() == v_q_id]
                
                st.markdown("---")
                st.markdown(f"<h2 style='color:#154360;'>📖 {v_q_id} {v_q_title}</h2>", unsafe_allow_html=True)
                
                records_to_print = []
                
                for idx, row in df_records.iterrows():
                    st.markdown(f"#### 📰 新聞標題：{row['新聞標題']}")
                    st.caption(f"📅 發布日期：{row['新聞日期']}")
                    
                    file_ids = str(row.get('照片清單', '')).split(',')
                    file_ids = [fid.strip() for fid in file_ids if fid.strip()]
                    
                    if file_ids:
                        table_html = "<table style='width:100%; table-layout:fixed; border-collapse:collapse; border:1px solid #D9E0E3; background-color:white; margin-bottom: 20px;'>"
                        for i in range(0, len(file_ids), 2):
                            table_html += "<tr>"
                            b64_img1 = get_drive_image_b64(file_ids[i])
                            if b64_img1:
                                table_html += f"<td style='border:1px solid #D9E0E3; padding:10px; text-align:center; width:50%; vertical-align:middle;'><img src='data:image/jpeg;base64,{b64_img1}' style='width:100%; max-height:300px; object-fit:contain; border-radius:8px;'></td>"
                            else:
                                table_html += "<td style='border:1px solid #D9E0E3; width:50%; text-align:center;'>圖片載入失敗</td>"
                                
                            if i + 1 < len(file_ids):
                                b64_img2 = get_drive_image_b64(file_ids[i+1])
                                if b64_img2:
                                    table_html += f"<td style='border:1px solid #D9E0E3; padding:10px; text-align:center; width:50%; vertical-align:middle;'><img src='data:image/jpeg;base64,{b64_img2}' style='width:100%; max-height:300px; object-fit:contain; border-radius:8px;'></td>"
                                else:
                                    table_html += "<td style='border:1px solid #D9E0E3; width:50%; text-align:center;'>圖片載入失敗</td>"
                            else:
                                table_html += "<td style='border:1px solid #D9E0E3; width:50%;'></td>"
                            table_html += "</tr>"
                        table_html += "</table>"
                        st.markdown(table_html, unsafe_allow_html=True)
                    
                    c_left, c_right = st.columns([1, 1])
                    with c_left:
                        st.markdown("**📄 原始新聞內容**")
                        st.markdown(f"<div style='background-color:#FDF6E3; padding:20px; border-radius:8px; border-left:5px solid #E6C27A; height:350px; overflow-y:auto; color:#555;'>{row['原始內容']}</div>", unsafe_allow_html=True)
                    with c_right:
                        st.markdown("**✨ Gemini 濃縮摘要 (含 SDGs)**")
                        fmt_ai = format_report_text_to_html(row['AI摘要'])
                        st.markdown(f"<div style='background-color:#E6F0F9; padding:20px; border-radius:8px; border-left:5px solid #8FAAB8; height:350px; overflow-y:auto; color:#154360; font-size:1.1em;'>{fmt_ai}</div>", unsafe_allow_html=True)
                    
                    st.markdown(f"<br>**🔗 原始新聞連結:** [{row['新聞連結']}]({row['新聞連結']})", unsafe_allow_html=True)
                    st.markdown("<hr style='border:1px dashed #ccc;'>", unsafe_allow_html=True)
                    
                    records_to_print.append({
                        '新聞標題': row['新聞標題'], '新聞日期': row['新聞日期'],
                        'AI摘要': row['AI摘要'], '照片清單': ",".join(file_ids),
                        '新聞連結': row['新聞連結']
                    })
                    
                st.markdown("<div style='text-align: center; margin-bottom: 10px;'><b>準備好下載這份 AI 智能彙整報告了嗎？</b></div>", unsafe_allow_html=True)
                with st.spinner("⏳ 產製 Word 檔中..."):
                    docx_bytes = generate_ai_word_report(v_q_id, v_q_title, records_to_print)
                    
                col1, col2, col3 = st.columns([3, 4, 3])
                with col2:
                    st.download_button(
                        label="📥 下載本題 AI 彙整報告 (Word檔)",
                        data=docx_bytes,
                        file_name=f"{v_q_id}_AI新聞彙整.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )