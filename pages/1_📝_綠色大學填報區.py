import streamlit as st
import pandas as pd
import gspread
import io
import re
import datetime
import base64
from PIL import Image
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn

# ==========================================
# 🌟 全域參數設定
# ==========================================
PDF_FILE_ID = "1hXs3yUwNxEiNZkJNeAfDTZjs9jxEgJos"
DRIVE_UPLOAD_FOLDER_ID = "1C1EbqH2oL-i_uqB6aEAw6S6tEwTvpnNc"

# ==========================================
# 🛡️ 資安防護罩
# ==========================================
if st.session_state.get("authentication_status") is not True:
    st.warning("⚠️ 請先至首頁登入系統！")
    st.stop()

st.set_page_config(page_title="嘉大綠色大學填報區", page_icon="📝", layout="wide")

# ==========================================
# 🎨 系統 UI 樣式設定
# ==========================================
st.markdown("""
<style>
    button[data-baseweb="tab"] { background-color: #E6F0F9 !important; border-radius: 8px 8px 0px 0px !important; margin-right: 4px !important; padding: 10px 20px !important; border: 1px solid #AED6F1 !important; border-bottom: none !important; transition: all 0.3s ease; }
    button[data-baseweb="tab"] p { font-size: 1.4em !important; font-weight: bold !important; color: #2C3E50 !important; }
    button[data-baseweb="tab"][aria-selected="true"] { background-color: #154360 !important; border: 1px solid #0B2331 !important; border-bottom: 3px solid #0B2331 !important; }
    button[data-baseweb="tab"][aria-selected="true"] p { color: #FFFFFF !important; }

    .morandi-select-title { background-color: #9DAB86; color: white; padding: 12px 15px; border-radius: 6px; font-weight: bold; font-size: 1.2em; margin-bottom: 5px; margin-top: 15px; }
    .morandi-question-title { background-color: #948B89; color: white; padding: 15px 18px; border-radius: 6px; font-weight: bold; font-size: 1.4em; margin-bottom: 15px; margin-top: 10px; }
    .morandi-orange-title { background-color: #D4A373; color: white; padding: 15px 20px; border-radius: 6px; font-weight: bold; font-size: 1.6em; margin-bottom: 15px; margin-top: 30px; text-align: center; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    .morandi-dark-title { background-color: #5C6B73; color: white; padding: 12px 20px; border-radius: 6px; font-weight: bold; font-size: 1.3em; margin-bottom: 10px; margin-top: 15px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
    
    .light-blue-box { background-color: #E6F0F9; color: #2C3E50; padding: 20px; border-left: 6px solid #8FAAB8; border-radius: 6px; margin-bottom: 15px; font-size: 1.05em; line-height: 1.8; }
    .light-yellow-box { background-color: #FDF6E3; color: #2C3E50; padding: 15px 20px; border-left: 6px solid #E6C27A; border-radius: 6px; margin-bottom: 5px; line-height: 1.6; }
    .light-brown-box { background-color: #F5EBE0; color: #2C3E50; padding: 15px 20px; border-left: 6px solid #D4A373; border-radius: 6px; margin-bottom: 5px; line-height: 1.6; }
    
    [data-testid="stExpander"] details summary { background-color: #2C3E50 !important; color: white !important; border-radius: 6px; }
    [data-testid="stExpander"] details summary p { color: white !important; font-size: 1.2em !important; font-weight: bold; }
    [data-testid="stExpander"] details summary svg { fill: white !important; }
    [data-testid="stExpander"] details[open] > div:nth-child(2) { background-color: #F4F6F7 !important; border: 2px solid #2C3E50; border-top: none; border-radius: 0 0 6px 6px; padding: 15px; }
    
    div.stButton > button[kind="secondary"] { background-color: #D6EAF8 !important; color: #154360 !important; border: 1px solid #AED6F1 !important; border-radius: 6px !important; font-weight: bold !important; }
    div.stButton > button[kind="primary"] { border-radius: 8px !important; font-weight: bold !important; font-size: 1.4em !important; padding: 12px 30px !important; }
    
    [data-testid="stDownloadButton"] button { background-color: #D4A373 !important; color: white !important; border: none !important; border-radius: 8px !important; font-weight: bold !important; font-size: 1.3em !important; padding: 12px 24px !important; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    [data-testid="stDownloadButton"] button:hover { background-color: #C39262 !important; }
    [data-testid="stDownloadButton"] button * { color: white !important; }

    [data-testid="stFileUploadDropzone"] small { display: none !important; }
    .custom-scrollbar::-webkit-scrollbar { width: 8px; }
    .custom-scrollbar::-webkit-scrollbar-track { background: #f1f1f1; border-radius: 4px; }
    .custom-scrollbar::-webkit-scrollbar-thumb { background: #8F9CA3; border-radius: 4px; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 🔌 初始化及資料庫連線功能
# ==========================================
@st.cache_resource
def get_gcp_credentials():
    skey = st.secrets["gcp_oauth"].to_dict()
    return Credentials(
        token=None, refresh_token=skey.get("refresh_token"), 
        token_uri="https://oauth2.googleapis.com/token", 
        client_id=skey.get("client_id"), client_secret=skey.get("client_secret")
    )

@st.cache_data(ttl=600)
def load_gsheet_data():
    try:
        creds = get_gcp_credentials()
        gc = gspread.authorize(creds)
        sh = gc.open_by_key('1JNbpZoZHWZRrIzn0whcQFnCDkOZghZmMyFidLE7dxT8')
        worksheet = sh.worksheet("評比題目表")
        data = worksheet.get_all_records()
        df = pd.DataFrame(data)
        df.columns = df.columns.str.strip()
        return df
    except Exception as e: return pd.DataFrame()

@st.cache_data(ttl=60)
def load_reported_data():
    try:
        creds = get_gcp_credentials()
        gc = gspread.authorize(creds)
        sh = gc.open_by_key('1JNbpZoZHWZRrIzn0whcQFnCDkOZghZmMyFidLE7dxT8')
        worksheet = sh.worksheet("填報資料庫")
        data = worksheet.get_all_records()
        return pd.DataFrame(data)
    except Exception as e: return pd.DataFrame()

# ==========================================
# 🌟 智慧上傳：照片無損壓縮引擎
# ==========================================
class CompressedFile:
    def __init__(self, buffer, name, mime_type):
        self.buffer = buffer
        self.name = name
        self.type = mime_type
    def getvalue(self): return self.buffer.getvalue()

def compress_image(file_obj):
    try:
        if not file_obj.type.startswith('image/'): return file_obj
        img = Image.open(file_obj)
        if img.format not in ['JPEG', 'PNG', 'JPG', 'WEBP']: return file_obj
        if img.mode in ("RGBA", "P"): img = img.convert("RGB")
        img.thumbnail((1920, 1920), Image.Resampling.LANCZOS)
        output = io.BytesIO()
        img.save(output, format='JPEG', quality=85, optimize=True)
        output.seek(0)
        new_name = file_obj.name.rsplit('.', 1)[0] + '.jpg'
        return CompressedFile(output, new_name, 'image/jpeg')
    except Exception as e: return file_obj

def upload_file_to_drive(file_obj, filename, folder_id):
    try:
        creds = get_gcp_credentials()
        drive_service = build('drive', 'v3', credentials=creds)
        file_metadata = {'name': filename, 'parents': [folder_id]}
        media = MediaIoBaseUpload(io.BytesIO(file_obj.getvalue()), mimetype=file_obj.type, resumable=True)
        uploaded_file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
        return uploaded_file.get('id')
    except Exception as e: return None

# ==========================================
# 🌟 智慧判斷檔案格式、照片方向與「全景圖偵測」
# ==========================================
@st.cache_data(ttl=3600, show_spinner=False)
def get_file_info(file_id, desc=""):
    try:
        creds = get_gcp_credentials()
        drive_service = build('drive', 'v3', credentials=creds)
        meta = drive_service.files().get(fileId=file_id, fields='mimeType').execute()
        mime_type = meta.get('mimeType', '')
        
        is_pdf = 'pdf' in mime_type.lower()
        is_image = mime_type.startswith('image/')
        
        request = drive_service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done: status, done = downloader.next_chunk()
        fh.seek(0)
        file_bytes = fh.read()
        
        is_landscape = True
        is_panorama = False # 新增全景圖偵測
        if is_image:
            try:
                img = Image.open(io.BytesIO(file_bytes))
                ratio = img.width / img.height
                if ratio < 1.0: is_landscape = False
                # 🌟 如果寬度是高度的 1.9 倍以上，判定為全景橫幅
                if ratio >= 1.9: is_panorama = True 
            except Exception: pass
            
        b64_str = base64.b64encode(file_bytes).decode('utf-8') if is_image else ""
        return {'id': file_id, 'desc': desc, 'is_pdf': is_pdf, 'is_image': is_image, 'is_landscape': is_landscape, 'is_panorama': is_panorama, 'b64': b64_str, 'mime_type': mime_type}
    except Exception as e: return None

def write_to_database(unit, reporter, ext, email, q_id, q_title, report_text, file_records):
    try:
        creds = get_gcp_credentials()
        gc = gspread.authorize(creds)
        sh = gc.open_by_key('1JNbpZoZHWZRrIzn0whcQFnCDkOZghZmMyFidLE7dxT8')
        ws = sh.worksheet("填報資料庫")
        now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        row = [now_str, unit, reporter, ext, email, q_id, q_title, report_text]
        for record in file_records[:10]: row.extend([record['desc'], record['id']])
        while len(row) < 28: row.append('')
        ws.append_row(row[:28])
        return True
    except Exception as e: 
        st.error(f"寫入資料庫失敗：{e}")
        return False

# ==========================================
# 🌟 公文級文字解析引擎：網頁完美縮排對齊
# ==========================================
def format_report_text_to_html(text):
    text = str(text)
    text = re.sub(r'(^|\n)(\d+[\.\)]|[-•*])\s*\n\s*', r'\1\2 ', text)
    html = ""
    for line in text.split('\n'):
        line = line.strip()
        if not line: 
            html += "<div style='height: 10px;'></div>"
            continue
        
        # 🌟 強化正則表達式，支援「1.」「2、」「(1)」「（一）」「-」「•」等多種公文常見標號
        match = re.match(r'^([0-9a-zA-Z]+[\.、\)]|[\(（][0-9a-zA-Z一二三四五六七八九十]+[\)）]|[一二三四五六七八九十]+、|[-•*])\s*(.*)', line)
        if match: 
            marker = match.group(1)
            content = match.group(2)
            # 設定左側預留空間，讓後續換行文字完美對齊
            html += f"<div style='display: flex; margin-bottom: 5px; line-height: 1.6;'><div style='width: 2.5em; flex-shrink: 0; text-align: right; padding-right: 0.5em;'>{marker}</div><div>{content}</div></div>"
        else: 
            html += f"<div style='margin-bottom: 5px; line-height: 1.6;'>{line}</div>"
    return html.replace('\n', ' ')

# ==========================================
# 🌟 Word 報告產製引擎 (公文級排版 + 全景圖支援)
# ==========================================
def set_run_font(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def generate_word_report(unit, reporter, ext, email, q_id, q_title, desc_text, req_text, report_text, enriched_files):
    doc = Document()
    doc.add_heading(f'嘉大綠色大學填報成果 - {unit}', 0)
    
    if reporter or ext or email:
        p_info = doc.add_paragraph(f"填報人：{reporter} | 分機：{ext} | Email：{email}")
        p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    doc.add_heading(f'{q_id} {q_title}', level=1)
    
    # 🌟 真正段落換行 (Enter)，避免版面跑掉
    doc.add_heading('題目說明：', level=2)
    for line in str(desc_text).replace('\r', '').split('\n'):
        if line.strip(): doc.add_paragraph(line.strip())
    
    doc.add_heading('資料需求：', level=2)
    clean_req = str(req_text).replace('<br>', '\n').replace('<b>', '').replace('</b>', '')
    for line in clean_req.replace('\r', '').split('\n'):
        if line.strip(): doc.add_paragraph(line.strip())
    
    doc.add_heading('填報資訊 / 年度執行亮點成果：', level=2)
    report_text_clean = re.sub(r'(^|\n)(\d+[\.\)]|[-•*])\s*\n\s*', r'\1\2 ', str(report_text))
    
    # 🌟 項目符號智慧懸掛縮排
    for line in report_text_clean.replace('\r', '').split('\n'):
        line = line.strip()
        if not line: continue
            
        p = doc.add_paragraph()
        match = re.match(r'^([0-9a-zA-Z]+[\.、\)]|[\(（][0-9a-zA-Z一二三四五六七八九十]+[\)）]|[一二三四五六七八九十]+、|[-•*])\s*(.*)', line)
        if match:
            marker = match.group(1)
            content = match.group(2)
            # 懸掛縮排設定：左側縮進 1 公分，第一行凸出 1 公分
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-1.0)
            # 使用 \t 定位點確保內容對齊
            p.add_run(f"{marker}\t{content}")
        else:
            p.add_run(line)
    
    doc.add_heading('佐證照片/檔案：', level=2)
    
    # 🌟 分離：全景圖、一般橫式、直式、文件
    panoramas = [f for f in enriched_files if f.get('is_image') and f.get('is_panorama')]
    landscapes = [f for f in enriched_files if f.get('is_image') and f.get('is_landscape') and not f.get('is_panorama')]
    portraits = [f for f in enriched_files if f.get('is_image') and not f.get('is_landscape')]
    other_docs = [f for f in enriched_files if not f.get('is_image')]
    
    # 1. 處理全景圖 (單獨佔據一整列，寬度加大)
    for pano in panoramas:
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            img_bytes = base64.b64decode(pano['b64'])
            fh = io.BytesIO(img_bytes)
            # 全景圖寬度設為 15 公分
            inline = p_img.add_run().add_picture(fh, width=Cm(15.0))
            try:
                pic_xml = inline._inline.xpath('.//pic:pic')[0]
                spPr = pic_xml.xpath('.//pic:spPr')[0]
                effectLst_xml = '<a:effectLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:softEdge rad="31750"/></a:effectLst>'
                spPr.append(parse_xml(effectLst_xml))
            except Exception: pass
        except Exception: p_img.add_run("(圖片無法插入)")
        p_text = cell.add_paragraph(f"{pano['desc']}")
        p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("") # 加空行區隔
    
    # 2. 處理一般配對照片
    pairs = []
    for i in range(0, len(landscapes) - 1, 2): pairs.append((landscapes[i], landscapes[i+1]))
    rem_l = landscapes[-1] if len(landscapes) % 2 != 0 else None
        
    for i in range(0, len(portraits) - 1, 2): pairs.append((portraits[i], portraits[i+1]))
    rem_p = portraits[-1] if len(portraits) % 2 != 0 else None
    
    if rem_l and rem_p: pairs.append((rem_l, rem_p)) 
    elif rem_l: pairs.append((rem_l, None))
    elif rem_p: pairs.append((rem_p, None))
    
    if len(pairs) > 0:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for p1, p2 in pairs:
            row = table.add_row()
            for idx, f in enumerate([p1, p2]):
                cell = row.cells[idx]
                if f:
                    p_img = cell.paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        img_bytes = base64.b64decode(f['b64'])
                        fh = io.BytesIO(img_bytes)
                        if f['is_landscape']: inline = p_img.add_run().add_picture(fh, width=Cm(7.5), height=Cm(5.5))
                        else: inline = p_img.add_run().add_picture(fh, width=Cm(7.0), height=Cm(9.0))
                        try:
                            pic_xml = inline._inline.xpath('.//pic:pic')[0]
                            spPr = pic_xml.xpath('.//pic:spPr')[0]
                            effectLst_xml = '<a:effectLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:softEdge rad="31750"/></a:effectLst>'
                            spPr.append(parse_xml(effectLst_xml))
                        except Exception: pass
                    except Exception: p_img.add_run("(圖片無法插入)")
                    p_text = cell.add_paragraph(f"{f['desc']}")
                    p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3. 處理文件
    if len(other_docs) > 0:
        for f in other_docs:
            p = doc.add_paragraph()
            p.add_run(f"📌 {f['desc']}").bold = True
            p.add_run(f"\n🔗 請至雲端檢視：https://drive.google.com/file/d/{f['id']}/view")
            
    if len(panoramas) == 0 and len(pairs) == 0 and len(other_docs) == 0: 
        doc.add_paragraph("此紀錄無上傳任何附件。")
    
    for p in doc.paragraphs:
        if not p.style.name.startswith('Heading'): p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs: set_run_font(r)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs: set_run_font(r)
        
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ==========================================
# 🚀 網頁主畫面與 UI 邏輯
# ==========================================
if "upload_count" not in st.session_state: 
    st.session_state.upload_count = 5
    
col_title, col_btn = st.columns([3, 1])
with col_title: st.title("📝 嘉大綠色大學填報系統")
with col_btn:
    st.write("") 
    if st.button("🔄 同步最新雲端資料", use_container_width=True):
        load_gsheet_data.clear()
        load_reported_data.clear()
        get_file_info.clear() 
        st.rerun()

tab_fill, tab_view = st.tabs(["📝 填報區", "📊 檢視填報成果"])

# ==========================================
# 🏷️ TAB 1: 填報區
# ==========================================
with tab_fill:
    if st.session_state.get('submit_success'):
        st.success("🎉 填報成功！資料已寫入資料庫，您可繼續填報其他項目！")
        st.session_state.submit_success = False

    with st.spinner('🔄 正在載入最新評比題目...'): df_questions = load_gsheet_data()

    if not df_questions.empty:
        unit_list = []
        for u in df_questions['權責單位'].dropna().unique():
            if str(u).strip() != '': unit_list.append(str(u).strip())
        
        c_unit, c_item = st.columns(2)
        with c_unit:
            st.markdown("<div class='morandi-select-title'>🏢 請選擇您的單位</div>", unsafe_allow_html=True)
            selected_unit = st.selectbox("", ["請選擇..."] + unit_list, label_visibility="collapsed")
        
        if selected_unit != "請選擇...":
            df_unit_questions = df_questions[df_questions['權責單位'].astype(str).str.strip() == selected_unit].copy()
            df_unit_questions['選項標示'] = df_unit_questions['當年度題目'].astype(str) + " - " + df_unit_questions['中文標題'].astype(str)
            
            with c_item:
                st.markdown("<div class='morandi-select-title'>📌 請選擇填報項目</div>", unsafe_allow_html=True)
                selected_option = st.selectbox("", ["請選擇..."] + df_unit_questions['選項標示'].tolist(), label_visibility="collapsed", key="sel_item")
            
            st.markdown("<div class='morandi-dark-title'>👤 填報人基本資料</div>", unsafe_allow_html=True)
            c_name, c_ext, c_email = st.columns(3)
            with c_name: reporter_name = st.text_input("填報人姓名", key="reporter_name", placeholder="請輸入姓名")
            with c_ext: reporter_ext = st.text_input("分機號碼", key="reporter_ext", placeholder="例如：1234")
            with c_email: reporter_email = st.text_input("電子郵件", key="reporter_email", placeholder="請輸入 Email")
            
            if selected_option != "請選擇...":
                selected_q_id = selected_option.split(" - ")[0]
                question_data = df_unit_questions[df_unit_questions['當年度題目'].astype(str) == selected_q_id].iloc[0]
                
                st.markdown("---")
                st.markdown(f"<div class='morandi-question-title'>📖 {question_data.get('當年度題目')}：{question_data.get('中文標題')}</div>", unsafe_allow_html=True)
                
                desc_text = str(question_data.get('中文說明', '無')).replace('中文說明：', '').replace('中文說明:', '').strip()
                st.markdown(f"<div style='color: #B85042; font-weight: bold; font-size: 1.3em; padding-left: 5px; margin-bottom: 25px;'>{desc_text}</div>", unsafe_allow_html=True)
                
                prev_q = str(question_data.get('前一年度題目', '')).strip()
                if not prev_q or prev_q.lower() in ['nan', '無', '']: 
                    st.info("💡 本項目為今年度新增，尚無去年提報參考資料。")
                else:
                    with st.expander("💡 點擊展開查看：前一年度參考資訊", expanded=True):
                        st.markdown(f"**對應之去年度題目：** {prev_q}\n---")
                        col_pdf, col_trans = st.columns([1, 1])
                        with col_pdf:
                            st.markdown("#### 📄 前一年度提報內容") 
                            pdf_id = str(question_data.get('單題PDF_ID', '')).strip()
                            if pdf_id and pdf_id.lower() != 'nan': 
                                st.markdown(f'<iframe src="https://drive.google.com/file/d/{pdf_id}/preview" width="100%" height="600" style="border: 2px solid #8F9CA3; border-radius: 8px;"></iframe>', unsafe_allow_html=True)
                            else: 
                                st.info("*(🚧 尚無專屬預覽檔案)*")
                        with col_trans:
                            st.markdown("#### 🇹🇼 中文翻譯參考")
                            ref_text = question_data.get('2025參考文字_AI預留', '')
                            if pd.notna(ref_text) and str(ref_text).strip() != "": 
                                st.markdown(f'<div class="custom-scrollbar" style="background-color: #FFFFFF; padding: 20px; border-radius: 8px; border: 2px solid #E2E7E3; height: 600px; overflow-y: auto; line-height: 1.8; font-size: 1.1em; color: #2C3E50; white-space: pre-wrap;">{ref_text}</div>', unsafe_allow_html=True)
                            else: 
                                st.info("*(🚧 尚無文字資料)*")
                
                st.markdown("<div class='morandi-orange-title'>資料填報區</div>", unsafe_allow_html=True)
                
                st.markdown("""
                <div class='light-blue-box'>
                    <strong style='font-size: 1.15em;'>提供資料及佐證示意照片說明：</strong><br>
                    <ul>
                        <li><b>成果說明文字：</b>請重點明確說明，包括質化、量化成果資訊。</li>
                        <li><b>佐證示意照片：</b>請提供 .jpg 檔，至少提供 3-4 張 (可多提供)，如說明有指定數量，請對應提供。</li>
                        <li>若需求資料僅有照片提供，則請您於「填報資訊 / 年度執行亮點成果」欄位填「<b>無須提供</b>」即可。</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
                
                req_text = str(question_data.get('資料需求', '無特別說明')).replace('\n', '<br>')
                st.markdown(f"""
                <div class='light-yellow-box'>
                    <strong style='font-size: 1.4em; color: #B85042;'>🔍 請貴單位協助提供下列資料：</strong><br>
                    <div style='font-size: 1.2em; margin-top: 10px; padding-left: 10px; color: #2C3E50;'>{req_text}</div>
                </div>
                """, unsafe_allow_html=True)
                
                st.markdown("<div class='morandi-dark-title'>✍️ 填報資訊 / 年度執行亮點成果</div>", unsafe_allow_html=True)
                report_text = st.text_area("填寫區", height=200, placeholder="請在此輸入您的填寫內容...", label_visibility="collapsed", key="report_input")
                
                st.markdown("""
                <div style='color: #555555; font-size: 1.0em; padding-left: 5px; margin-top: 5px; margin-bottom: 25px; line-height: 1.6;'>
                    <b>📝 編輯小提示：</b><br>
                    ．可直接使用數字 (1. 2.) 或連字號 (-) 來列出項目符號。<br>
                    ．可於右下角手動調整輸入框高度，以利檢視。
                </div>
                """, unsafe_allow_html=True)
                
                st.markdown("<div class='morandi-dark-title'>📎 請上傳照片或檔案</div>", unsafe_allow_html=True)
                st.markdown("<div style='color: #2C3E50; font-weight:bold; margin-top:10px; margin-bottom:10px;'>請先設定上傳檔案數量：</div>", unsafe_allow_html=True)
                
                col_add, col_sub, _ = st.columns([2, 2, 6])
                if col_add.button("➕ 新增一筆檔案區"): 
                    if st.session_state.upload_count < 10:
                        st.session_state.upload_count += 1
                        st.rerun()
                    else: st.warning("最多僅支援上傳 10 組檔案！")
                        
                if col_sub.button("➖ 減少一筆檔案區"):
                    if st.session_state.upload_count > 1: 
                        st.session_state.upload_count -= 1
                        st.rerun()
                
                st.write("")
                for i in range(st.session_state.upload_count):
                    c1, c2 = st.columns([1, 1])
                    with c1:
                        st.markdown(f"<div style='font-size:1.1em; font-weight:bold; margin-bottom:5px; margin-top:5px;'>📌 檔案 {i+1} 資料說明：</div>", unsafe_allow_html=True)
                        st.text_input("說明", key=f"desc_{i}", placeholder="例如：智慧路燈系統畫面", label_visibility="collapsed")
                    with c2:
                        st.markdown(f"<div style='margin-bottom:5px; margin-top:5px;'>&nbsp;</div>", unsafe_allow_html=True)
                        st.file_uploader("選擇檔案", key=f"file_{i}", label_visibility="collapsed")
                
                st.markdown("<br><div style='display: flex; justify-content: center;'>", unsafe_allow_html=True)
                submitted = st.button("📤 資料確認送出", type="primary", use_container_width=True)
                st.markdown("</div>", unsafe_allow_html=True)
                
                if submitted:
                    valid_files = []
                    for i in range(st.session_state.upload_count):
                        f_obj = st.session_state.get(f"file_{i}")
                        if f_obj is not None:
                            desc = st.session_state.get(f"desc_{i}", f"未命名附件{i+1}")
                            valid_files.append({"file": f_obj, "desc": desc})
                            
                    if not report_text.strip() and len(valid_files) == 0: 
                        st.error("⚠️ 請至少填寫成果說明，或上傳一份佐證檔案！")
                    else:
                        with st.spinner("⏳ 正在無損壓縮檔案並寫入資料庫..."):
                            upload_records = []
                            for i, vf in enumerate(valid_files):
                                optimized_file = compress_image(vf['file'])
                                ext = optimized_file.name.split('.')[-1]
                                safe_desc = vf['desc'].strip() if vf['desc'].strip() else f"未命名附件{i+1}"
                                new_filename = f"{selected_q_id}-{safe_desc}.{ext}"
                                
                                file_id = upload_file_to_drive(optimized_file, new_filename, DRIVE_UPLOAD_FOLDER_ID)
                                if file_id: upload_records.append({'desc': safe_desc, 'id': file_id})
                            
                            db_success = write_to_database(
                                selected_unit, reporter_name, reporter_ext, reporter_email, 
                                selected_q_id, question_data.get('中文標題', ''), report_text, upload_records
                            )
                            
                            if db_success:
                                st.session_state.submit_success = True
                                keys_to_delete = []
                                for key in list(st.session_state.keys()):
                                    if key in ["sel_item", "report_input"] or key.startswith("desc_") or key.startswith("file_"): 
                                        keys_to_delete.append(key)
                                for key in keys_to_delete: del st.session_state[key]
                                st.session_state.upload_count = 5
                                load_reported_data.clear() 
                                st.rerun()

# ==========================================
# 🏷️ TAB 2: 檢視填報成果 (左右 4:6 黃金排版)
# ==========================================
with tab_view:
    df_reported = load_reported_data()
    
    if df_reported.empty: 
        st.info("💡 目前資料庫中尚未有任何填報紀錄。")
    else:
        st.markdown("""
        <div style='background-color: #FFF3E0; color: #E67E22; padding: 10px 15px; border-left: 5px solid #E67E22; border-radius: 5px; margin-bottom: 20px; font-weight: bold;'>
            💡 發現填報資料有誤嗎？請直接回到「📝 填報區」重新填寫並送出，系統會自動以您「最新一次」的送出紀錄為準喔！
        </div>
        """, unsafe_allow_html=True)

        rep_units = []
        for u in df_reported['權責單位'].dropna().unique():
            if str(u).strip() != '': rep_units.append(str(u).strip())
        
        c_v_unit, c_v_item = st.columns(2)
        with c_v_unit:
            st.markdown("<div class='morandi-select-title'>🏢 請選擇欲查詢之單位</div>", unsafe_allow_html=True)
            view_unit = st.selectbox("", ["請選擇..."] + rep_units, label_visibility="collapsed", key="v_unit")
            
        if view_unit != "請選擇...":
            df_view_unit = df_reported[df_reported['權責單位'].astype(str).str.strip() == view_unit]
            
            rep_items = []
            for item in (df_view_unit['題號'].astype(str) + " - " + df_view_unit['中文標題'].astype(str)).unique():
                rep_items.append(item)
            
            with c_v_item:
                st.markdown("<div class='morandi-select-title'>📌 請選擇已填報之項目</div>", unsafe_allow_html=True)
                view_item = st.selectbox("", ["請選擇..."] + rep_items, label_visibility="collapsed", key="v_item")
                
            if view_item != "請选择...":
                v_q_id = view_item.split(" - ")[0]
                latest_record = df_view_unit[df_view_unit['題號'].astype(str) == v_q_id].iloc[-1]
                
                v_q_title = latest_record.get('中文標題', '')
                v_desc_text = "無"
                v_req_text = "無特別說明"
                
                if 'df_questions' in locals() and not df_questions.empty:
                    q_match = df_questions[df_questions['當年度題目'].astype(str) == v_q_id]
                    if not q_match.empty:
                        v_desc_text = str(q_match.iloc[0].get('中文說明', '無')).replace('中文說明：', '').strip()
                        v_req_text = str(q_match.iloc[0].get('資料需求', '無特別說明'))
                
                st.markdown("---")
                
                # 🌟 填報人資訊深灰色字體、放大一號
                v_reporter = latest_record.get('填報人', '無')
                v_ext = latest_record.get('填報人分機', '無')
                v_email = latest_record.get('填報人電子郵件', '無')
                st.markdown(f"<div style='color: #555555; font-size: 1.15em; margin-bottom: 10px; font-weight: bold;'>👤 填報人：{v_reporter} &nbsp;|&nbsp; 📞 分機：{v_ext} &nbsp;|&nbsp; 📧 Email：{v_email}</div>", unsafe_allow_html=True)
                
                st.markdown(f"<div class='morandi-question-title'>📖 {v_q_id}：{v_q_title}</div>", unsafe_allow_html=True)
                st.markdown(f"<div style='color: #B85042; font-weight: bold; font-size: 1.3em; padding-left: 5px; margin-bottom: 25px;'>{v_desc_text}</div>", unsafe_allow_html=True)
                
                st.markdown(f"""
                <div class='light-brown-box'>
                    <strong style='font-size: 1.4em; color: #B85042;'>🔍 資料需求：</strong><br>
                    <div style='font-size: 1.2em; margin-top: 10px; padding-left: 10px; color: #2C3E50;'>{v_req_text.replace(chr(10), '<br>')}</div>
                </div>
                """, unsafe_allow_html=True)
                
                # 🌟 版面切分為左右 4:6 黃金比例
                col_left, col_right = st.columns([4, 6])
                
                # ==========================
                # 左邊 (40%)：填報資訊區
                # ==========================
                with col_left:
                    st.markdown("<div class='morandi-dark-title'>✍️ 填報資訊 / 年度執行亮點成果</div>", unsafe_allow_html=True)
                    formatted_report = format_report_text_to_html(latest_record.get('填報內容', '無'))
                    st.markdown(f"<div style='background-color: #F8FAFB; padding: 20px; border-radius: 8px; border: 1px solid #E2E7E3; font-size: 1.1em; color: #2C3E50; margin-bottom: 25px;'>{formatted_report}</div>", unsafe_allow_html=True)
                
                # ==========================
                # 右邊 (60%)：佐證照片或檔案區
                # ==========================
                with col_right:
                    st.markdown("<div class='morandi-dark-title'>📎 佐證照片或檔案</div>", unsafe_allow_html=True)
                    
                    raw_files = []
                    for i in range(1, 11):
                        desc = str(latest_record.get(f'檔案{i}_說明', '')).strip()
                        fid = str(latest_record.get(f'檔案{i}_ID', '')).strip()
                        if fid: raw_files.append({'desc': desc, 'id': fid})
                            
                    enriched_files = []
                    for rf in raw_files:
                        info = get_file_info(rf['id'], rf['desc'])
                        if info: enriched_files.append(info)
                    
                    panoramas = [f for f in enriched_files if f.get('is_image') and f.get('is_panorama')]
                    landscapes = [f for f in enriched_files if f.get('is_image') and f.get('is_landscape') and not f.get('is_panorama')]
                    portraits = [f for f in enriched_files if f.get('is_image') and not f.get('is_landscape')]
                    other_docs = [f for f in enriched_files if not f.get('is_image')]
                    
                    if len(enriched_files) == 0:
                        st.info("此紀錄無上傳任何附件。")
                    else:
                        # 🌟 1. 處理全景圖 (單獨顯示，不切割)
                        for pano in panoramas:
                            st.markdown(f"**📌 {pano['desc']}**")
                            img_html = f"""
                            <div style="text-align: center;">
                                <img src="data:{pano['mime_type']};base64,{pano['b64']}" 
                                     style="width: 100%; height: auto; object-fit: contain; background-color: #f1f1f1; border-radius: 8px; margin-bottom: 20px; border: 1px solid #ccc; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                            </div>
                            """
                            st.markdown(img_html, unsafe_allow_html=True)

                        # 🌟 2. 處理一般照片 (並排顯示，取消 cover 裁切，保證完整呈現)
                        pairs = []
                        for i in range(0, len(landscapes) - 1, 2): pairs.append((landscapes[i], landscapes[i+1]))
                        rem_l = landscapes[-1] if len(landscapes) % 2 != 0 else None
                        for i in range(0, len(portraits) - 1, 2): pairs.append((portraits[i], portraits[i+1]))
                        rem_p = portraits[-1] if len(portraits) % 2 != 0 else None
                        
                        if rem_l and rem_p: pairs.append((rem_l, rem_p))
                        elif rem_l: pairs.append((rem_l, None))
                        elif rem_p: pairs.append((rem_p, None))
                        
                        for p1, p2 in pairs:
                            cols = st.columns(2)
                            for idx, f in enumerate([p1, p2]):
                                if f:
                                    with cols[idx]:
                                        st.markdown(f"**📌 {f['desc']}**")
                                        img_html = f"""
                                        <div style="text-align: center;">
                                            <img src="data:{f['mime_type']};base64,{f['b64']}" 
                                                 style="width: 100%; height: auto; object-fit: contain; background-color: #f1f1f1; border-radius: 8px; margin-bottom: 20px; border: 1px solid #ccc; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                                        </div>
                                        """
                                        st.markdown(img_html, unsafe_allow_html=True)
                        
                        # 🌟 3. 處理其他文件
                        if len(other_docs) > 0:
                            for f in other_docs:
                                drive_link = f"https://drive.google.com/file/d/{f['id']}/view"
                                icon, ftype = ("📄", "PDF 檔案") if f.get('is_pdf') else ("📁", "其他文件")
                                st.markdown(f"""
                                <div style='background-color: #f1f1f1; border-radius: 8px; border: 1px solid #ccc; padding: 25px; text-align: center; margin-bottom: 20px;'>
                                    <span style='font-size:1.5em;'>{icon}</span> <span style='font-size:1.1em; font-weight:bold;'>📌 {f['desc']} ({ftype})</span><br><br>
                                    <a href='{drive_link}' target='_blank' style='background-color: #5C6B73; color: white; padding: 10px 15px; border-radius: 6px; text-decoration: none; font-weight: bold;'>🔗 點擊前往雲端檢視</a>
                                </div>
                                """, unsafe_allow_html=True)
                
                st.markdown("---")
                st.markdown("<div style='text-align: center; margin-bottom: 10px;'><b>想要留存這份填報紀錄嗎？</b></div>", unsafe_allow_html=True)
                
                with st.spinner("⏳ 正在為您產生 Word 報告檔案..."):
                    docx_bytes = generate_word_report(
                        view_unit, v_reporter, v_ext, v_email, v_q_id, v_q_title, 
                        v_desc_text, v_req_text, latest_record.get('填報內容', ''), enriched_files
                    )
                
                col_dl1, col_dl2, col_dl3 = st.columns([3, 4, 3])
                with col_dl2:
                    st.download_button(
                        label="📥 下載填報成果 (Word檔)",
                        data=docx_bytes,
                        file_name=f"{view_unit}_{v_q_id}_填報成果.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )