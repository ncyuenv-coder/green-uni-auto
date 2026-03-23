import streamlit as st
import pandas as pd
import gspread
import io
import re
import datetime
import base64
import zipfile
from PIL import Image
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import qn

# ==========================================
# 🌟 全域參數設定 & 資安防護罩
# ==========================================
# [精準導入 3] 狀態管理防護：確保驗證狀態變數存在
if st.session_state.get("authentication_status") is not True:
    st.warning("⚠️ 請先至首頁登入系統！")
    st.stop()

if st.session_state.get("username") != "admin_ui":
    st.error("🚫 權限不足！此頁面僅限系統管理員 (admin_ui) 存取。")
    st.stop()

st.set_page_config(page_title="嘉大綠色大學填報管理區", page_icon="📊", layout="wide")

# ==========================================
# 🎨 系統 UI 樣式設定
# ==========================================
st.markdown("""
<style>
    button[data-baseweb="tab"] { 
        background-color: #E6F0F9 !important; border-radius: 8px 8px 0px 0px !important; 
        margin-right: 4px !important; padding: 10px 20px !important; 
        border: 1px solid #AED6F1 !important; border-bottom: none !important; transition: all 0.3s ease; 
    }
    button[data-baseweb="tab"] p { font-size: 1.4em !important; font-weight: bold !important; color: #2C3E50 !important; }
    button[data-baseweb="tab"][aria-selected="true"] { 
        background-color: #154360 !important; border: 1px solid #0B2331 !important; border-bottom: 3px solid #0B2331 !important; 
    }
    button[data-baseweb="tab"][aria-selected="true"] p { color: #FFFFFF !important; }

    .stRadio div[role="radiogroup"] { flex-direction: row !important; gap: 15px !important; margin-bottom: 20px !important; }
    .stRadio div[role="radiogroup"] label {
        background-color: #F0F3F4 !important; border: 2px solid #BDC3C7 !important; border-radius: 10px !important; 
        padding: 12px 25px !important; cursor: pointer; transition: all 0.2s;
    }
    .stRadio div[role="radiogroup"] label:hover { background-color: #E5E8E8 !important; }
    .stRadio div[role="radiogroup"] label p { font-size: 1.2rem !important; font-weight: 800 !important; color: #566573 !important; margin: 0 !important; }
    .stRadio div[role="radiogroup"] label[data-checked="true"] { background-color: #34495E !important; border-color: #34495E !important; }
    .stRadio div[role="radiogroup"] label[data-checked="true"] p { color: #FFFFFF !important; }

    .morandi-select-title { background-color: #9DAB86; color: white; padding: 12px 15px; border-radius: 6px; font-weight: bold; font-size: 1.2em; margin-bottom: 5px; margin-top: 15px; }
    .morandi-question-title { background-color: #948B89; color: white; padding: 15px 18px; border-radius: 6px; font-weight: bold; font-size: 1.4em; margin-bottom: 15px; margin-top: 10px; }
    .morandi-dark-title { background-color: #5C6B73; color: white; padding: 12px 20px; border-radius: 6px; font-weight: bold; font-size: 1.3em; margin-bottom: 10px; margin-top: 15px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
    .light-brown-box { background-color: #F5EBE0; color: #2C3E50; padding: 15px 20px; border-left: 6px solid #D4A373; border-radius: 6px; margin-bottom: 5px; line-height: 1.6; }
    
    div.stButton > button[kind="primary"] { border-radius: 8px !important; font-weight: bold !important; font-size: 1.4em !important; padding: 12px 30px !important; }
    [data-testid="stDownloadButton"] button { background-color: #D4A373 !important; color: white !important; border: none !important; border-radius: 8px !important; font-weight: bold !important; font-size: 1.3em !important; padding: 12px 24px !important; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    [data-testid="stDownloadButton"] button:hover { background-color: #C39262 !important; }
    [data-testid="stDownloadButton"] button * { color: white !important; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 🔌 初始化及資料庫連線功能
# ==========================================
# [精準導入 2] 快取資源：確保 Google API 連線憑證不重複建立
@st.cache_resource
def get_gcp_credentials():
    skey = st.secrets["gcp_oauth"].to_dict()
    return Credentials(
        token=None, refresh_token=skey.get("refresh_token"), 
        token_uri="https://oauth2.googleapis.com/token", 
        client_id=skey.get("client_id"), client_secret=skey.get("client_secret")
    )

# [精準導入 2] 快取資料：將 ttl 延長至 86400，減輕 API 負擔，並交由 UI 按鈕主動更新
@st.cache_data(ttl=86400)
def load_gsheet_data(sheet_name):
    try:
        creds = get_gcp_credentials()
        gc = gspread.authorize(creds)
        sh = gc.open_by_key('1JNbpZoZHWZRrIzn0whcQFnCDkOZghZmMyFidLE7dxT8')
        df = pd.DataFrame(sh.worksheet(sheet_name).get_all_records())
        df.columns = df.columns.str.strip()
        return df
    except Exception: 
        return pd.DataFrame()

# [精準導入 2] 快取資料：雲端圖片資訊快取
@st.cache_data(ttl=3600, show_spinner=False)
def get_file_info(file_id, desc=""):
    try:
        creds = get_gcp_credentials()
        drive_service = build('drive', 'v3', credentials=creds)
        meta = drive_service.files().get(fileId=file_id, fields='mimeType').execute()
        mime_type = meta.get('mimeType', '')
        
        is_pdf = 'pdf' in mime_type.lower()
        is_image = mime_type.startswith('image/')
        
        request = drive_service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done: 
            status, done = downloader.next_chunk()
        fh.seek(0)
        file_bytes = fh.read()
        
        is_landscape = True
        is_panorama = False
        if is_image:
            try:
                img = Image.open(io.BytesIO(file_bytes))
                ratio = img.width / img.height
                if ratio < 1.0: is_landscape = False
                if ratio >= 1.9: is_panorama = True 
            except Exception: pass
            
        b64_str = base64.b64encode(file_bytes).decode('utf-8') if is_image else ""
        return {
            'id': file_id, 'desc': desc, 'is_pdf': is_pdf, 'is_image': is_image, 
            'is_landscape': is_landscape, 'is_panorama': is_panorama, 
            'b64': b64_str, 'mime_type': mime_type
        }
    except Exception: 
        return None

# ==========================================
# 🌟 排版模組 (HTML & Word 表格共用引擎)
# ==========================================
def format_report_text_to_html(text):
    text = str(text)
    text = re.sub(r'(^|\n)(\d+[\.\)]|[-•*])\s*\n\s*', r'\1\2 ', text)
    html = ""
    for line in text.split('\n'):
        line = line.strip()
        if not line: 
            html += "<div style='height: 10px;'></div>"
            continue
        if re.match(r'^【.*】$', line):
            html += f"<div style='font-size: 1.15em; font-weight: bold; color: #154360; margin-top: 15px; margin-bottom: 5px; border-bottom: 2px solid #AED6F1; padding-bottom: 3px;'>{line}</div>"
            continue
        match = re.match(r'^([0-9a-zA-Z]+[\.、\)]|[\(（][0-9a-zA-Z一二三四五六七八九十]+[\)）]|[一二三四五六七八九十]+、|[-•*])\s*(.*)', line)
        if match: 
            marker = match.group(1); content = match.group(2)
            html += f"<div style='padding-left: 2.2em; text-indent: -2.2em; margin-bottom: 5px; line-height: 1.6;'>{marker} {content}</div>"
        else: 
            html += f"<div style='margin-bottom: 5px; line-height: 1.6;'>{line}</div>"
    return html.replace('\n', ' ')

def generate_html_image_table(enriched_files_dict):
    """UI 呈現：4張照片為1列，超寬照片跨2欄，下方維持放置照片說明文字"""
    has_image = any(f.get('is_image') for files in enriched_files_dict.values() for f in files)
    if not has_image: return ""
    
    table_html = "<table style='width:100%; table-layout:fixed; border-collapse:collapse; border:1px solid #D9E0E3; background-color:white; margin-bottom: 25px;'>"
    for section_title, files in enriched_files_dict.items():
        images = [f for f in files if f.get('is_image')]
        if not images: continue
        
        if len(enriched_files_dict) > 1:
            table_html += f"<tr><td colspan='4' style='border:1px solid #D9E0E3; padding:10px; text-align:center; background-color:#E6F0F9;'><strong style='font-size:1.2em; color:#154360;'>{section_title}</strong></td></tr>"
        
        table_html += "<tr>"
        slots_filled = 0
        
        for f in images:
            slots_needed = 2 if f.get('is_panorama') else 1
            if slots_filled + slots_needed > 4:
                for _ in range(4 - slots_filled):
                    table_html += "<td style='border:1px solid #D9E0E3; width:25%;'></td>"
                table_html += "</tr><tr>"
                slots_filled = 0
            
            width_pct = 50 if slots_needed == 2 else 25
            img_height = "auto" if slots_needed == 2 else "200px"
            table_html += f"<td colspan='{slots_needed}' style='border:1px solid #D9E0E3; padding:15px; text-align:center; vertical-align:top; width:{width_pct}%;'><img src='data:{f['mime_type']};base64,{f['b64']}' style='width:100%; height:{img_height}; object-fit:contain; background-color:#f1f1f1; border-radius:8px; margin-bottom:10px;'><br><b>{f['desc']}</b></td>"
            slots_filled += slots_needed

        if 0 < slots_filled < 4:
            for _ in range(4 - slots_filled):
                table_html += "<td style='border:1px solid #D9E0E3; width:25%;'></td>"
            table_html += "</tr>"
            
    table_html += "</table>"
    return table_html

def add_images_to_word_table(doc, enriched_files_dict):
    """Word 產製：照片與文字分離成2列，超寬照片跨欄置中，同類照片優先並排，嚴格控制長寬尺寸"""
    has_image = any(f.get('is_image') for files in enriched_files_dict.values() for f in files)
    if not has_image: return
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    for section_title, files in enriched_files_dict.items():
        images = [f for f in files if f.get('is_image')]
        if not images: continue
        
        if len(enriched_files_dict) > 1:
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[1])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(section_title).bold = True
            
        panoramas = [f for f in images if f.get('is_panorama')]
        landscapes = [f for f in images if f.get('is_landscape') and not f.get('is_panorama')]
        portraits = [f for f in images if not f.get('is_landscape')]
        
        chunks = []
        for p_img in panoramas:
            chunks.append(([p_img], 2))
        for i in range(0, len(landscapes) - 1, 2):
            chunks.append((landscapes[i:i+2], 1))
        rem_l = landscapes[-1] if len(landscapes) % 2 != 0 else None
        for i in range(0, len(portraits) - 1, 2):
            chunks.append((portraits[i:i+2], 1))
        rem_p = portraits[-1] if len(portraits) % 2 != 0 else None
        if rem_l and rem_p:
            chunks.append(([rem_l, rem_p], 1))
        elif rem_l:
            chunks.append(([rem_l], 1))
        elif rem_p:
            chunks.append(([rem_p], 1))

        for chunk_files, slots_type in chunks:
            row_img = table.add_row()
            row_desc = table.add_row()
            if slots_type == 2:
                f = chunk_files[0]
                cell_img = row_img.cells[0]
                cell_img.merge(row_img.cells[1])
                cell_desc = row_desc.cells[0]
                cell_desc.merge(row_desc.cells[1])
                cell_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p_img = cell_img.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    img_bytes = base64.b64decode(f['b64'])
                    inline = p_img.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(15.5))
                except: p_img.add_run("(圖片無法插入)")
                cell_desc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p_desc = cell_desc.paragraphs[0]
                p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_desc.add_run(f"{f['desc']}")
            else:
                for j in range(2):
                    cell_img = row_img.cells[j]
                    cell_desc = row_desc.cells[j]
                    cell_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell_desc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if j < len(chunk_files):
                        f = chunk_files[j]
                        p_img = cell_img.paragraphs[0]
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        try:
                            img_bytes = base64.b64decode(f['b64'])
                            if f.get('is_landscape'):
                                p_img.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(7.5), height=Cm(5.5))
                            else:
                                p_img.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(7.0), height=Cm(9.0))
                        except: p_img.add_run("(圖片無法插入)")
                        p_desc = cell_desc.paragraphs[0]
                        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_desc.add_run(f"{f['desc']}")

def set_run_font(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

# ==========================================
# 🌟 核心資料處理與 Word 產製流程
# ==========================================
def extract_single_question_data(v_q_id, df_q, df_r):
    q_match_base = df_q[df_q['當年度題目'].astype(str).str.strip() == v_q_id]
    v_q_title = str(q_match_base.iloc[0]['中文標題']).strip() if not q_match_base.empty else "未知標題"
    v_desc_text = str(q_match_base.iloc[0].get('中文說明', '無')).replace('中文說明：', '').strip() if not q_match_base.empty else "無"
    
    df_q_records = df_r[df_r['題號'].astype(str).str.strip() == v_q_id].copy()
    if not df_q_records.empty:
        df_q_records['填報時間'] = pd.to_datetime(df_q_records['填報時間'])
        df_q_records = df_q_records.sort_values('填報時間').drop_duplicates(subset=['權責單位'], keep='last')
    
    info_strings_ui = []
    info_strings_word = [] 
    combined_reqs = ""
    combined_texts = ""
    unit_files_dict = {}
    unit_count = len(df_q_records)
    
    for idx, row in df_q_records.iterrows():
        u = str(row['權責單位']).strip()
        rep = str(row.get('填報人', '無')).strip()
        ext = str(row.get('填報人分機', '無')).strip()
        em = str(row.get('填報人電子郵件', '無')).strip()
        
        info_strings_ui.append(f"填報單位：{u}  |  👤 填報人：{rep}  |  📞 分機：{ext}  |  📧 Email：{em}")
        info_strings_word.append(f"填報單位：{u}  |  填報人：{rep}  |  分機：{ext}  |  Email：{em}")
        
        q_match_unit = df_q[(df_q['當年度題目'].astype(str).str.strip() == v_q_id) & (df_q['權責單位'].astype(str).str.strip() == u)]
        req_text = str(q_match_unit.iloc[0].get('資料需求', '無特別說明')) if not q_match_unit.empty else (str(q_match_base.iloc[0].get('資料需求', '無')) if not q_match_base.empty else "無")
        rep_text = str(row.get('填報內容', '無')).strip()
        
        if unit_count > 1:
            combined_reqs += f"【{u}】\n{req_text}\n\n"
            combined_texts += f"【{u}】\n{rep_text}\n\n"
        else:
            combined_reqs = req_text; combined_texts = rep_text
            
        files_for_u = []
        for i in range(1, 11):
            raw_id = str(row.get(f'檔案{i}_ID', '')).strip()
            desc = str(row.get(f'檔案{i}_說明', '')).strip()
            if raw_id:
                id_match = re.search(r'[-\w]{25,}', raw_id)
                if id_match:
                    files_for_u.append({'id': id_match.group(0), 'desc': desc})
        unit_files_dict[u] = files_for_u

    unit_enriched_files = {}
    for u, files in unit_files_dict.items():
        enriched = [get_file_info(f['id'], f['desc']) for f in files]
        unit_enriched_files[f"【{u}】"] = [e for e in enriched if e]
        
    return v_q_title, v_desc_text, info_strings_ui, info_strings_word, combined_reqs.strip(), combined_texts.strip(), unit_enriched_files

def build_word_document(v_q_id, v_q_title, info_strings_word, desc_text, req_text, report_text, unit_enriched_files):
    doc = Document()
    doc.add_heading(f'嘉大綠色大學填報成果彙整', 0)
    for info in info_strings_word: 
        doc.add_paragraph(info).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    doc.add_heading(f'{v_q_id} {v_q_title}', level=1)
    
    doc.add_heading('題目說明：', level=2)
    for line in str(desc_text).replace('\r', '').split('\n'):
        if line.strip(): doc.add_paragraph(line.strip())
        
    doc.add_heading('資料需求：', level=2)
    for line in str(req_text).replace('<br>', '\n').replace('<b>', '').replace('</b>', '').replace('\r', '').split('\n'):
        if line.strip(): doc.add_paragraph(line.strip())
        
    doc.add_heading('填報資訊 / 年度執行亮點成果：', level=2)
    report_text_clean = re.sub(r'(^|\n)(\d+[\.\)]|[-•*])\s*\n\s*', r'\1\2 ', str(report_text))
    for line in report_text_clean.replace('\r', '').split('\n'):
        line = line.strip()
        if not line: continue
        p = doc.add_paragraph()
        if re.match(r'^【.*】$', line): p.add_run(line).bold = True
        else:
            match = re.match(r'^([0-9a-zA-Z]+[\.、\)]|[\(（][0-9a-zA-Z一二三四五六七八九十]+[\)）]|[一二三四五六七八九十]+、|[-•*])\s*(.*)', line)
            if match:
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.first_line_indent = Cm(-0.8)
                p.add_run(f"{match.group(1)} {match.group(2)}")
            else: p.add_run(line)
            
    doc.add_heading('佐證照片/檔案：', level=2)
    add_images_to_word_table(doc, unit_enriched_files)
    
    has_any_doc = any(not f.get('is_image') for files in unit_enriched_files.values() for f in files)
    if has_any_doc:
        doc.add_heading('其他參考資料：', level=2)
        for u, files in unit_enriched_files.items():
            docs = [f for f in files if not f.get('is_image')]
            if docs:
                if len(unit_enriched_files) > 1: doc.add_paragraph(f"{u}").runs[0].bold = True
                for f in docs:
                    doc.add_paragraph().add_run(f"{f['desc']} (此為 {'PDF 檔案' if f.get('is_pdf') else '其他文件'})").bold = True
                    doc.add_paragraph(f"請至雲端檢視：https://drive.google.com/file/d/{f['id']}/view")
    
    if not any(f.get('is_image') for files in unit_enriched_files.values() for f in files) and not has_any_doc:
        doc.add_paragraph("此紀錄無上傳任何單位附件。")
    
    for p in doc.paragraphs:
        if not p.style.name.startswith('Heading'): p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs: set_run_font(r)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs: set_run_font(r)
                    
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ==========================================
# [精準導入 1] 分頁模組：局部重跑 Fragment
# ==========================================
@st.fragment
def render_tab_missing():
    with st.spinner("⏳ 正在比對資料庫，找出未申報名單..."):
        df_req = load_gsheet_data("評比題目表")
        df_sub = load_gsheet_data("填報資料庫")
        
        if df_req.empty:
            st.error("無法讀取評比題目表，請確認連線。")
        else:
            df_req_clean = df_req[['權責單位', '當年度題目', '中文標題']].copy()
            df_req_clean['權責單位'] = df_req_clean['權責單位'].astype(str).str.strip()
            df_req_clean['當年度題目'] = df_req_clean['當年度題目'].astype(str).str.strip()
            df_req_clean = df_req_clean[df_req_clean['權責單位'] != ''].drop_duplicates()
            
            df_sub_clean = pd.DataFrame(columns=['權責單位', '當年度題目'])
            if not df_sub.empty:
                df_sub_clean = df_sub[['權責單位', '題號']].copy()
                df_sub_clean.rename(columns={'題號': '當年度題目'}, inplace=True)
                df_sub_clean['權責單位'] = df_sub_clean['權責單位'].astype(str).str.strip()
                df_sub_clean['當年度題目'] = df_sub_clean['當年度題目'].astype(str).str.strip()
                df_sub_clean = df_sub_clean.drop_duplicates()
            
            merged = df_req_clean.merge(df_sub_clean, on=['權責單位', '當年度題目'], how='left', indicator=True)
            missing = merged[merged['_merge'] == 'left_only'].drop(columns=['_merge'])
            
            st.markdown("<div class='morandi-dark-title'>🔍 尚未填報之單位與題目清單</div>", unsafe_allow_html=True)
            
            if missing.empty:
                st.success("🎉 太棒了！所有單位皆已完成所有題目的填報！")
            else:
                st.info(f"💡 目前共有 **{len(missing)}** 筆待填報項目。")
                grouped_missing = missing.groupby('權責單位')
                
                html_str = ""
                count = 0
                for unit, group in grouped_missing:
                    bg_summary = "#9CB4C4" if count % 2 == 0 else "#D0DCE3"
                    bg_content = "#E8EEF2" if count % 2 == 0 else "#F7F9FA"
                    border_color = bg_summary
                        
                    html_str += f'''
                    <details style="margin-bottom: 12px; border: 1px solid {border_color}; border-radius: 8px; box-shadow: 0 2px 5px rgba(0,0,0,0.05);">
                        <summary style="background-color: {bg_summary}; color: black; padding: 15px 20px; font-size: 1.25em; font-weight: bold; border-radius: 8px; cursor: pointer;">
                            🏢 {unit} <span style="color: #C0392B; font-size: 0.9em;">(尚缺 {len(group)} 題)</span>
                        </summary>
                        <div style="background-color: {bg_content}; padding: 20px; color: black; border-radius: 0 0 8px 8px; border-top: 1px solid {border_color};">
                            <ul style="margin: 0; padding-left: 25px; line-height: 1.8; font-size: 1.15em;">
                    '''
                    for idx, row in group.iterrows():
                        html_str += f"<li style='margin-bottom: 8px;'><b>{row['當年度題目']}</b>：{row['中文標題']}</li>"
                    html_str += "</ul></div></details>"
                    count += 1
                st.markdown(html_str, unsafe_allow_html=True)

@st.fragment
def render_tab_download():
    df_r = load_gsheet_data("填報資料庫")
    df_q = load_gsheet_data("評比題目表")
    
    if df_r.empty: 
        st.info("💡 目前資料庫中尚未有任何填報紀錄可供下載。")
    else:
        st.markdown("<br>", unsafe_allow_html=True)
        dl_mode = st.radio("選擇下載模式", ["單題彙整下載", "全部題目彙整下載"], horizontal=True, label_visibility="collapsed")
        
        reported_q_ids = df_r['題號'].astype(str).str.strip().unique().tolist()
        reported_q_ids.sort()
        
        if dl_mode == "單題彙整下載":
            st.markdown("<div class='morandi-select-title'>📌 請選擇欲彙整下載之題目</div>", unsafe_allow_html=True)
            q_options = []
            for qid in reported_q_ids:
                title_match = df_q[df_q['當年度題目'].astype(str).str.strip() == qid]
                title = title_match.iloc[0]['中文標題'] if not title_match.empty else "未知標題"
                q_options.append(f"{qid} - {title}")
                
            q_options.sort()
            view_item = st.selectbox("", ["請選擇..."] + q_options, label_visibility="collapsed", key="m_item")
            
            if view_item != "請選擇...":
                v_q_id = view_item.split(" - ")[0]
                
                with st.spinner("⏳ 正在拉取資料與雲端圖片，請稍候..."):
                    v_q_title, v_desc_text, info_strings_ui, info_strings_word, combined_reqs, combined_texts, unit_enriched_files = extract_single_question_data(v_q_id, df_q, df_r)
                
                st.markdown("---")
                for info in info_strings_ui:
                    st.markdown(f"<div style='color: #555555; font-size: 1.15em; margin-bottom: 5px; font-weight: bold;'>{info}</div>", unsafe_allow_html=True)
                
                st.markdown(f"<div class='morandi-question-title'>📖 {v_q_id}：{v_q_title}</div>", unsafe_allow_html=True)
                st.markdown(f"<div style='color: #B85042; font-weight: bold; font-size: 1.3em; padding-left: 5px; margin-bottom: 25px;'>{v_desc_text}</div>", unsafe_allow_html=True)
                
                display_reqs = combined_reqs.replace('\n', '<br>')
                display_reqs = re.sub(r'【(.*?)】', r'<span style="color:#154360; font-weight:bold;">【\1】</span>', display_reqs)
                st.markdown(f"""
                <div class='light-brown-box'>
                    <strong style='font-size: 1.4em; color: #B85042;'>🔍 資料需求：</strong><br>
                    <div style='font-size: 1.2em; margin-top: 10px; padding-left: 10px; color: #2C3E50;'>{display_reqs}</div>
                </div>
                """, unsafe_allow_html=True)
                
                st.markdown("<div class='morandi-dark-title'>✍️ 填報資訊 / 年度執行亮點成果</div>", unsafe_allow_html=True)
                formatted_report = format_report_text_to_html(combined_texts)
                st.markdown(f"<div style='background-color: #F8FAFB; padding: 20px; border-radius: 8px; border: 1px solid #E2E7E3; font-size: 1.1em; color: #2C3E50; margin-bottom: 25px;'>{formatted_report}</div>", unsafe_allow_html=True)
                
                st.markdown("<div class='morandi-dark-title'>📎 佐證照片或檔案</div>", unsafe_allow_html=True)
                unit_html_table = generate_html_image_table(unit_enriched_files)
                if unit_html_table: st.markdown(unit_html_table, unsafe_allow_html=True)
                else: st.info("此紀錄無上傳任何單位附件。")
                
                st.markdown("<div style='text-align: center; margin-top: 30px; margin-bottom: 10px;'><b>準備好彙整這份報告了嗎？</b></div>", unsafe_allow_html=True)
                
                with st.spinner("⏳ 正在為您產生 Word 報告檔案..."):
                    docx_bytes = build_word_document(v_q_id, v_q_title, info_strings_word, v_desc_text, combined_reqs, combined_texts, unit_enriched_files)
                
                col_dl1, col_dl2, col_dl3 = st.columns([3, 4, 3])
                with col_dl2:
                    st.download_button(
                        label="📥 下載單題彙整成果 (Word檔)",
                        data=docx_bytes,
                        file_name=f"{v_q_id}_填報彙整成果.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    
        elif dl_mode == "全部題目彙整下載":
            st.markdown("<div class='morandi-select-title'>📦 批量打包下載所有題目報告</div>", unsafe_allow_html=True)
            st.info("💡 此功能將自動為「所有已填報之題目」分別產生獨立的 Word 報告（含對應照片佐證），並打包成一個 ZIP 壓縮檔，方便您一次性下載留存。")
            
            if st.button("🚀 一鍵打包產生 ZIP 壓縮檔", type="primary"):
                with st.spinner("⏳ 正在背景逐題解析資料、抓取圖片並產生 Word 報告，這需要幾分鐘的時間，請耐心等候..."):
                    zip_buffer = io.BytesIO()
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                        total_q = len(reported_q_ids)
                        for i, qid in enumerate(reported_q_ids):
                            status_text.text(f"正在處理第 {i+1}/{total_q} 題：{qid} ...")
                            
                            v_q_title, v_desc_text, info_strings_ui, info_strings_word, combined_reqs, combined_texts, unit_enriched_files = extract_single_question_data(qid, df_q, df_r)
                            docx_bytes = build_word_document(qid, v_q_title, info_strings_word, v_desc_text, combined_reqs, combined_texts, unit_enriched_files)
                            
                            safe_title = re.sub(r'[\\/*?:"<>|]', "", v_q_title) 
                            file_name = f"{qid}_{safe_title}_填報彙整.docx"
                            zip_file.writestr(file_name, docx_bytes.getvalue())
                            
                            progress_bar.progress((i + 1) / total_q)
                            
                    zip_buffer.seek(0)
                    status_text.empty()
                    st.success("✅ 全部打包完成！請點擊下方按鈕下載。")
                    
                    st.download_button(
                        label="📥 下載全題彙整成果 (ZIP壓縮檔)",
                        data=zip_buffer,
                        file_name=f"嘉大綠色大學_全題彙整_{datetime.datetime.now().strftime('%Y%m%d')}.zip",
                        mime="application/zip",
                        use_container_width=True
                    )

# ==========================================
# 🚀 網頁主畫面與 UI 邏輯
# ==========================================
col_title, col_btn = st.columns([3, 1])
with col_title: st.title("📊 嘉大綠色大學填報管理區")
with col_btn:
    st.write("") 
    if st.button("🔄 同步最新雲端資料", use_container_width=True):
        load_gsheet_data.clear()
        get_file_info.clear() 
        st.rerun()

tab_missing, tab_download = st.tabs(["🎯 追蹤未申報單位", "📥 下載彙整資料"])

with tab_missing:
    render_tab_missing()

with tab_download:
    render_tab_download()